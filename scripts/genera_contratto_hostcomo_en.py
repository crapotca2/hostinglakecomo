# -*- coding: utf-8 -*-
# Rebuild the Host Como EN management agreement from the extracted text and
# apply the 8 client-requested modifications, then render a styled .docx.
import re, sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\Andrei\.claude\projects\c--Users-Andrei-Desktop-Claudio-Hosting-Lake-Como\b1cd929f-e070-4d43-9b6a-76b4e56974f7\tool-results\b76izosjp.txt"
OUT = r"C:\Users\Andrei\Desktop\contratti borghi\Contratto Host Como EN - rev1.docx"

# ---------------- parse ----------------
def clean_line(l):
    l = l.replace("<sub>2</sub>", "\u2082").replace("<sub>", "").replace("</sub>", "")
    for a, b in [("\u2019", "'"), ("\u2018", "'"), ("\u201c", '"'), ("\u201d", '"'), ("\u00a0", " ")]:
        l = l.replace(a, b)
    return l.rstrip()

raw = open(SRC, encoding="cp1252", errors="replace").read().split("\n")
lines = []
for l in raw[2:]:
    s = l.strip()
    if not s:
        lines.append(""); continue
    if re.match(r"^HOST COMO . Mandate and Services Agreement$", s): continue
    if re.match(r"^Pag\.\s*\d+$", s): continue
    lines.append(clean_line(l))

SUBCLAUSE = r"^\d{1,2}\.\d{1,2}(-bis|-ter)?\."
def classify(t):
    t = t.strip()
    if not t: return None
    if re.match(r"^\d{1,2}\.\s+[A-Z]", t) and not re.match(r"^\d{1,2}\.\d", t): return "article"
    if re.match(SUBCLAUSE, t): return "subclause"
    if re.match(r"^[a-z]\)\s", t): return "letter"
    if re.match(r"^\u2022", t): return "bullet"
    if re.match(r"^[A-D]\.\s+[A-Z]", t): return "recital"
    if re.match(r"^SCHEDULE\s+[A-D]", t): return "schedule"
    if re.match(r"^Section\s+[IVX]+", t): return "section"
    if re.match(r"^[A-D]\.(I{1,2}\.)?\d", t): return "sched_item"
    for p in [r"^WHEREAS$", r"^NOW, THEREFORE", r"^Place and date", r"^_____",
              r"^The Principal", r"^Angelo Talarico \(the Agent\)", r"^Specific approval of onerous",
              r"^\(tick the relevant box\)", r"^\[\s?\]", r"^Principal's obligations connected",
              r"^MANDATE AND SERVICES AGREEMENT", r"^Courtesy English translation",
              r"^This mandate and services agreement", r"^References to legislative sources",
              r"^the Principal and the Agent hereinafter"]:
        if re.match(p, t): return "special"
    return None

REF_TAIL = re.compile(r"(art\.|artt\.|para\.|paragraph|no\.|Decree-Law|Law)\s*$", re.I)
blocks, cur, in_defs = [], "", False
for l in lines:
    if l == "": continue
    t = l.strip()
    if re.match(r"^1\.2\.", t): in_defs = True
    if re.match(r"^1\.3\.", t) or re.match(r"^2\.\s+[A-Z]", t): in_defs = False
    kind = classify(t)
    if in_defs and re.match(r'^"', t): kind = "definition"
    if kind == "subclause" and REF_TAIL.search(cur): kind = None
    if kind or not cur:
        if cur: blocks.append(cur.strip())
        cur = l
    else:
        cur += " " + t
if cur: blocks.append(cur.strip())
print("parsed blocks:", len(blocks))

# ---------------- modification helpers ----------------
WARN = []
def _find(sub):
    idx = [i for i, b in enumerate(blocks) if sub in b]
    return idx
def sub(locator, old, new, n=1):
    ix = _find(locator)
    if not ix: WARN.append(f"LOCATOR NOT FOUND: {locator[:60]!r}"); return
    i = ix[0]
    if old not in blocks[i]: WARN.append(f"OLD NOT FOUND in block {i}: {old[:60]!r}"); return
    blocks[i] = blocks[i].replace(old, new, n)
def replace_block(locator, new_text):
    ix = _find(locator)
    if not ix: WARN.append(f"BLOCK NOT FOUND: {locator[:60]!r}"); return
    blocks[ix[0]] = new_text
def delete_block(locator):
    ix = _find(locator)
    if not ix: WARN.append(f"DELETE NOT FOUND: {locator[:60]!r}"); return
    blocks.pop(ix[0])
def insert_after(locator, new_blocks):
    ix = _find(locator)
    if not ix: WARN.append(f"INSERT-AFTER NOT FOUND: {locator[:60]!r}"); return
    i = ix[0] + 1
    for nb in new_blocks:
        blocks.insert(i, nb); i += 1

# ---------------- MOD 1: tourist tax monthly ----------------
sub("summary of the tourist taxes collected", "on a quarterly basis", "on a monthly basis")
sub("remitting it to the Principal on a quarterly basis together with a detailed statement",
    "on a quarterly basis together with a detailed statement", "on a monthly basis together with a detailed statement")
sub("quarterly delivery to the Principal with a detailed statement",
    "quarterly delivery to the Principal", "monthly delivery to the Principal")
sub("The following are included in the 15% commission referred to in art. 7.1",
    "with quarterly delivery to the Principal", "with monthly delivery to the Principal")

# ---------------- MOD 2: insurance recommendation (no HC PI) ----------------
sub("A copy of the policy in force shall be produced to the Agent upon written request within 15 (fifteen) days of such request;",
    "days of such request;",
    "days of such request. The Agent, without acting as an insurance intermediary, recommends that the Principal take out a dedicated short-term-rental / holiday-home insurance policy and may indicate suitable products, including online solutions, for the Principal to evaluate independently;")

# ---------------- MOD 3: direct bookings 3 tiers ----------------
NEW_71BIS = ("7.1-bis. Direct bookings, owner-sourced sales and free stays. Bookings not sourced through the Agent's "
"channels are treated as follows. (a) Owner-sourced paid bookings (the Principal's own sales). Where the Principal "
"itself sources a paid booking (by way of example, a paying acquaintance or friend of the Principal) and the Agent "
"is involved by providing guest communication and assistance, the Fee is reduced to 5% (five per cent) of the gross "
"consideration agreed with the Guest, plus VAT where applicable, in consideration of the communication and "
"guest-assistance services rendered by the Agent; the mandatory Alloggiati Web/CIN registration and the handling of "
"the tourist tax for such bookings are carried out by the Agent and are included in that reduced Fee. Where instead "
"the Agent renders the full range of Services on such a booking, the ordinary Fee of 15% under art. 7.1 applies. "
"(b) Free stays (no Fee). No Fee is due for stays granted free of charge, without any direct or indirect "
"consideration, by the Principal to itself, to its family members (the spouse, the civil partner under Law no. 76 of "
"20/5/2016, the cohabiting partner or party to a registered de facto cohabitation under the same Law 76/2016, "
"children, parents and siblings) or to friends and other persons hosted by the Principal free of charge, provided "
"that such stays are, in the aggregate: (i) notified to the Agent in writing with not less than 15 (fifteen) days' "
"notice before check-in, for the removal of availability on the channel manager and for the Alloggiati Web/CIN "
"filings; (ii) for aggregate periods not exceeding 30 (thirty) calendar days per calendar year; and (iii) without "
"payment of any direct or indirect consideration. Failure to give the prior notice required above for an "
"owner-sourced paid booking or for a free stay entails application of the Fee, at the rate applicable under this "
"article, on the average consideration of the last 12 months for an equivalent overnight stay, without prejudice to "
"further contractual remedies in the event of repeated breach.")
replace_block("7.1-bis. Direct bookings and personal use by the Principal.", NEW_71BIS)

sub("a Fee of 15% (fifteen per cent) remains due to the Agent, calculated",
    "a Fee of 15% (fifteen per cent) remains due to the Agent, calculated",
    "a Fee of 5% (five per cent) remains due to the Agent, calculated")

# vessatorie descriptions
sub("art. 7.1-bis (application of the Fee only to direct bookings",
    "art. 7.1-bis (application of the Fee only to direct bookings for which the Agent actually renders the Services; exemption for bookings managed independently by the Principal; exceptions for free personal stays of the Principal and of family members within the limits of 30 days/year and 15 days' notice)",
    "art. 7.1-bis (15% Fee where the Agent renders the full Services; reduced 5% Fee on owner-sourced paid bookings where the Agent provides communication and guest assistance; no Fee on free stays of the Principal, family members and friends within the limits of 30 days/year and 15 days' notice)")
sub("Fee of 15% always due on the net consideration actually collected",
    "Fee of 15% always due on the net consideration actually collected",
    "reduced Fee of 5% due on the net consideration actually collected")

# ---------------- MOD 5: balanced liability ----------------
sub("d) the provision of services by third-party undertakings (cleaning company, laundry, maintenance technicians) where the Agent has acted with the required diligence",
    "d) the provision of services by third-party undertakings (cleaning company, laundry, maintenance technicians) where the Agent has acted with the required diligence in their selection and coordination;",
    "d) damage from the provision of services by third-party undertakings (cleaning company, laundry, maintenance technicians) only to the extent it results from causes not attributable to the Agent's selection, coordination or supervision, it being understood that the Agent is otherwise responsible for its operators and subcontractors pursuant to art. 1228 of the Italian Civil Code as set out in art. 6.2-bis;")
sub("e) breach, delay in performance or non-performance of its obligations due to causes beyond its control",
    "platform overbooking,", "platform overbooking not attributable to the Agent's management of the channel manager,")
sub("Liability towards the Principal for the acts of such parties remains with the Agent within the limits of art. 6.",
    "Liability towards the Principal for the acts of such parties remains with the Agent within the limits of art. 6.",
    "The Agent is responsible towards the Principal for the acts and omissions of such parties pursuant to art. 1228 of the Italian Civil Code, within the limits of art. 6 and as set out in art. 6.2-bis.")
NEW_62BIS = ("6.2-bis. Responsibility of the Agent for subcontractors, calendar, overbooking and filings. Without prejudice "
"to art. 6.1 as amended and to the cap in art. 6.3, the Agent is responsible towards the Principal for: (i) the acts and "
"omissions of its own team and of the cleaning operators, suppliers and subcontractors that it selects and remunerates, "
"pursuant to art. 1228 of the Italian Civil Code; (ii) calendar "
"errors and overbookings arising from its own management of the channel manager and of the OTA calendars; and (iii) the "
"penalties, interest and charges arising from the late, incomplete or omitted performance, through its own fault, of the "
"regulatory filings it has undertaken under art. 3.4 (Alloggiati Web registration, monthly ISTAT filings, and collection, "
"tracking and remittance of the tourist tax). The responsibility under point (i) does not extend to the interventions of "
"specialised technicians (plumbers, electricians, masons and the like) whose performance and cost the Principal authorises "
"and bears under art. 3.2(h) and Schedule A.I.4: for such interventions the Agent is responsible only for their diligent "
"selection and coordination on the Principal's instruction. This responsibility does not extend to facts attributable to "
"the Principal's sphere of control (art. 3.2), to information or credentials not timely provided by the Principal, or to "
"events of force majeure under art. 6.1(e), and operates within the limits of art. 6.3 save in the cases in which that cap "
"does not apply.")
insert_after("6.2. The Agent is economically liable for damage caused directly by it or by its operators/subcontractors",
             [NEW_62BIS])
# A.I.4: masons added + owner keeps go-ahead and cost of specialist interventions
sub("A.I.4. Coordination of maintenance and reporting of damage.",
    "specialised technicians (plumbers, electricians, aerial installers); the costs of spare parts, materials and technicians' services are borne by the Principal and invoiced separately, with the Principal's prior approval for work exceeding Euro 100.00 (one hundred/00);",
    "specialised technicians (plumbers, electricians, masons, aerial installers); the costs of spare parts, materials and technicians' services are borne by the Principal and invoiced separately, with the Principal's prior approval for work exceeding Euro 100.00 (one hundred/00). Save for the urgent minor works under point (ii) below, the decision whether to carry out, and the cost of, any intervention by specialised technicians (plumbers, electricians, masons and the like) rest entirely with the Principal, the Agent's role being limited to the diligent selection and coordination of such technicians on the Principal's instruction;")

# ---------------- MOD 6: termination symmetric + cure + remove EUR200 ----------------
delete_block("8.2-bis. Liquidated damages for the Principal's early withdrawal within 6 months.")
sub("the Principal may withdraw from this Agreement upon 30 (thirty) days' notice, without payment of any liquidated damages and without application of art. 8.2-bis.",
    "without payment of any liquidated damages and without application of art. 8.2-bis.",
    "without payment of any amount by way of liquidated damages.")
sub("the Principal has a corresponding right of withdrawal, without application of the liquidated damages referred to in art. 8.2-bis even where exercised within 6 months of Publication of the Listing, with notice reduced to 30 (thirty) days",
    "without application of the liquidated damages referred to in art. 8.2-bis even where exercised within 6 months of Publication of the Listing, ",
    "")
sub("either Party may terminate the Agreement pursuant to art. 1456 of the Italian Civil Code without notice and without application of the liquidated damages referred to in art. 8.2-bis.",
    "and without application of the liquidated damages referred to in art. 8.2-bis.",
    "and without any liquidated damages.")
sub("The Principal retains the right to withdraw from the Agreement, without notice and without application of the liquidated damages referred to in art. 8.2-bis, should it not accept the assignment",
    "without application of the liquidated damages referred to in art. 8.2-bis",
    "without any liquidated damages")
sub("Conversely, the Principal may terminate the Agreement pursuant to art. 1456 of the Italian Civil Code in the event of repeated and material breach by the Agent of the obligations under art. 3.4.",
    "Conversely, the Principal may terminate the Agreement pursuant to art. 1456 of the Italian Civil Code in the event of repeated and material breach by the Agent of the obligations under art. 3.4.",
    "Conversely, the Principal may terminate the Agreement pursuant to art. 1456 of the Italian Civil Code in the event of material breach by the Agent of the obligations under artt. 3.4, 6.2-bis and 8.3(a). The express termination remedies under this article are subject to the cure procedure in art. 8.4-bis.")
NEW_84BIS = ("8.4-bis. Cure period and reciprocity. Save for the cases of automatic suspension or of termination without a "
"remedy period expressly provided for elsewhere in this Agreement (in particular art. 3.6 on the CIN), where a breach is "
"remediable the Party intending to terminate shall first serve on the other a written notice to perform (diffida ad "
"adempiere pursuant to art. 1454 of the Italian Civil Code) granting a period of not less than 15 (fifteen) days to remedy; "
"the Agreement shall be terminated only if the breach is not remedied within that period. The express termination remedy "
"under art. 1456 of the Italian Civil Code operates on a reciprocal basis for the material breaches expressly identified in "
"this Agreement, both for the Agent (art. 8.4) and for the Principal (breach by the Agent of artt. 3.4, 6.2-bis and 8.3(a)).")
insert_after("8.4. The Agent may terminate this Agreement pursuant to art. 1456", [NEW_84BIS])
# vessatorie cleanups
sub("art. 8.2-bis (liquidated damages of Euro 200.00 for the Principal's early withdrawal within 6 months, with exclusions);",
    "art. 8.2-bis (liquidated damages of Euro 200.00 for the Principal's early withdrawal within 6 months, with exclusions); ", "")
sub("art. 8.4 (express termination clauses);",
    "art. 8.4 (express termination clauses);", "artt. 8.4 and 8.4-bis (reciprocal express termination clauses, with a 15-day cure period);")

# ---------------- MOD 7: Property definition multi-unit ----------------
sub('"Property": has the meaning set out in Recital A;',
    '"Property": has the meaning set out in Recital A;',
    '"Property": has the meaning set out in Recital A; where more than one property is entrusted to the Agent, the term "Property" refers to each of the properties listed in Schedule E, and the obligations, the Fee and the availability threshold under this Agreement apply separately to each such property;')

# ---------------- MOD 4 + 8 + 7: new schedules/clauses via sentinels ----------------
insert_after("C.7. The cleaning company engaged does not provide", [
    "C.8. Price list for cleaning, linen, check-in and additional services. The following are the indicative rates for the "
    "operational and additional services; the cleaning and linen costs are borne by the Principal and passed through at the "
    "supplier's rates, and any additional Service under art. 2.2 is agreed in writing in advance. The cleaning service is "
    "carried out by the team selected and remunerated by the Agent, at a cost per turnover of between Euro 50.00 and Euro "
    "100.00 depending on the property; the linen and laundry service is provided by the same cleaning company and is charged "
    "at cost. Interventions by specialised technicians (plumbers, electricians, masons) are quoted "
    "case by case, require the Principal's prior go-ahead and are borne by the Principal (art. 3.2(h) and Schedule A.I.4). "
    "The remaining figures below are to be completed and confirmed by the Agent before Publication of the Listing.",
    "[[TABLE:pricelist]]",
    "C.9. Illustrative total cost (cost stack). The following non-binding example illustrates, on a nightly gross "
    "consideration of Euro 100.00, the combined effect of the OTA commission (borne by the Principal), the Host Como Fee of "
    "15% and VAT (where applicable under the Agent's tax regime), down to the amount net to the Principal, before the "
    "cedolare secca (21% on the first property and 26% from the second) which the Principal settles separately.",
    "[[TABLE:coststack]]",
])
insert_after("D.6. Liability. Each Party shall be liable for damage arising from the processing of personal data", [
    "SCHEDULE E \u2014 THE PROPERTIES",
    "This Schedule identifies each Property entrusted to the Agent. Where more than one Property is entrusted, a separate "
    "annex in the form below is completed for each unit; the obligations, the Fee and the 150-day availability threshold "
    "apply separately to each Property.",
    "[[TABLE:properties]]",
    "Note: duplicate this Schedule E annex for each additional Property.",
])

if WARN:
    print("\n*** WARNINGS ***")
    for w in WARN: print("  -", w)
else:
    print("all modifications applied cleanly.")

# ---------------- render docx ----------------
doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)

def bold_leading(p, text):
    m = re.match(r"^((?:\d{1,2}\.\d{1,2}(?:-bis|-ter)?\.)|(?:\d{1,2}\.)|(?:[A-D]\.(?:I{1,2}\.)?\d+\.)|(?:[a-z]\)))\s+", text)
    if m:
        r = p.add_run(m.group(1) + " "); r.bold = True
        p.add_run(text[m.end():])
    else:
        p.add_run(text)

PRICELIST = [("Service", "Unit", "Rate EUR (excl. VAT)")] + [
    ("Standard cleaning", "per turnover", "50\u2013100 (by property)"),
    ("Linen & laundry", "provided by the cleaning company", "at cost"),
    ("Check-in outside standard hours / holidays", "per check-in", "[\u2026]"),
    ("Additional multilingual check-in", "per check-in", "[\u2026]"),
    ("Extraordinary / deep cleaning", "per intervention", "[\u2026]"),
    ("Smart-lock installation & management", "one-off / monthly", "[\u2026]"),
    ("Personalised welcome basket", "per booking", "[\u2026]"),
    ("Translation / personalised guide", "per item", "[\u2026]"),
    ("Restyling / home staging / photo shoot", "on quote", "[\u2026]"),
]
COSTSTACK = [("Item", "Amount")] + [
    ("Gross nightly consideration", "EUR 100.00"),
    ("\u2212 OTA commission (Airbnb / Booking.com / Expedia, ~[\u2026]%) \u2013 borne by the Principal", "\u2212 EUR [\u2026]"),
    ("\u2212 Host Como Fee 15% (on the gross)", "\u2212 EUR 15.00"),
    ("\u2212 VAT on the Host Como Fee ([\u2026]% per the Agent's regime; nil if flat-rate)", "\u2212 EUR [\u2026]"),
    ("= Amount net to the Principal (before cedolare secca)", "EUR [\u2026]"),
]
PROPERTIES = [("Field", "Property 1", "Property 2")] + [
    ("Address", "[\u2026]", "[\u2026]"),
    ("Cadastral details (foglio / part. / sub.)", "[\u2026]", "[\u2026]"),
    ("CIN", "[\u2026]", "[\u2026]"),
    ("CIR (if applicable)", "[\u2026]", "[\u2026]"),
    ("Person capacity", "[\u2026]", "[\u2026]"),
    ("Costs (utilities / IMU / condo, notes)", "[\u2026]", "[\u2026]"),
    ("Equipment / notes", "[\u2026]", "[\u2026]"),
]
def add_table(data):
    t = doc.add_table(rows=len(data), cols=len(data[0])); t.style = "Table Grid"
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            c = t.cell(ri, ci); c.text = str(val)
            for par in c.paragraphs:
                for rn in par.runs:
                    rn.font.size = Pt(9)
                    if ri == 0: rn.bold = True

for b in blocks:
    if b.startswith("[[TABLE:"):
        add_table({"[[TABLE:pricelist]]": PRICELIST, "[[TABLE:coststack]]": COSTSTACK, "[[TABLE:properties]]": PROPERTIES}[b])
        continue
    k = classify(b)
    if b == blocks[0]:  # title
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(b); r.bold = True; r.font.size = Pt(13); continue
    if b.startswith("Courtesy English translation"):
        p = doc.add_paragraph(); r = p.add_run(b); r.italic = True; r.font.size = Pt(9); continue
    if k == "article" or b.startswith("SCHEDULE "):
        # split schedule heading from any inline body
        m = re.match(r"^(SCHEDULE [A-D][^a-z]+?)\s+([A-Z][a-z].*)$", b)
        head, body = (m.group(1), m.group(2)) if m else (b, None)
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
        r = p.add_run(head); r.bold = True; r.font.size = Pt(12)
        if body:
            pb = doc.add_paragraph(); pb.add_run(body)
        continue
    if k == "section":
        p = doc.add_paragraph(); r = p.add_run(b); r.bold = True; r.italic = True; continue
    if k == "special" and (b.startswith("WHEREAS") or b.startswith("NOW, THEREFORE")):
        p = doc.add_paragraph(); r = p.add_run(b); r.bold = True; continue
    if b.startswith("Specific approval of onerous clauses"):
        m = re.match(r"^(Specific approval of onerous clauses \([^)]*\))\s+(.*)$", b)
        if m:
            p = doc.add_paragraph(); r = p.add_run(m.group(1)); r.bold = True
            doc.add_paragraph(m.group(2)); continue
    p = doc.add_paragraph()
    if k in ("letter", "bullet"):
        p.paragraph_format.left_indent = Inches(0.3)
    if k == "bullet":
        p.add_run(b)
    else:
        bold_leading(p, b)

doc.save(OUT)
print("SAVED", OUT)
print("total rendered blocks:", len(blocks))
