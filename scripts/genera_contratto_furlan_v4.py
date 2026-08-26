"""
Genera il contratto Host Como v4 per Andrea Furlan + Ephanie Sacramento
sull'immobile Via della Libertà 7, Cernobbio.

Versione minimale: aderente al PDF originale, con sole modifiche
strettamente necessarie per la copropriétà:
 - Intestazione con entrambi i Mandanti
 - Plurale "i Mandanti" dove serve grammaticalmente
 - Angelo Talarico con dati completi (indirizzo Como)
 - Firma di Andrea + Ephanie in 3 sedi
 - Luogo e data firma lasciati in bianco
Nessuna clausola aggiuntiva (no Premessa A-bis, no art. 4.4 IBAN unico,
no rewrite fiscale svizzero, no con-Titolari GDPR, no allegato D adeg.).
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

# --- DATI COMPILATI ---
MANDANTE1 = {
    "nome": "ANDREA FURLAN",
    "nato_luogo": "Treviso",
    "nato_data": "14/01/1973",
    "residenza": "Via della Libertà 7, 22012 Cernobbio (CO)",
    "cf": "FRLNDR73A14L407J",
}
MANDANTE2 = {
    "nome": "EPHANIE SACRAMENTO",
    "nato_luogo": "Amneville (Francia)",
    "nato_data": "29/01/1981",
    "residenza": "Chemin de l'Asile 5, 3960 Sierre (Svizzera)",
}
MANDATARIO = {
    "nome": "ANGELO TALARICO",
    "nato_luogo": "Filadelfia (VV)",
    "nato_data": "24/07/1963",
    "cf": "TLRNGL63L24C933I",
    "domicilio": "Via Nino Bixio 4, 22100 Como (CO)",
    "piva": "04157560139",
    "ateco": "68.32.00",
    "pec": "angelo.talarico@mpspec.it",
    "marchio": "HOST COMO",
    "sito": "hostcomo.com",
}
IMMOBILE = {
    "indirizzo": "Via della Libertà 7, 22012 Cernobbio (CO)",
    "comune": "Cernobbio",
    "sezione_urbana": "ROV",
    "foglio": "10",
    "particella": "38",
    "subalterno": "701",
    "piani": "Piano Terra + Piano Primo + Piano Secondo",
}

OUTPUT = Path(r"C:/Users/Andrei/Desktop/Contratto_HostComo_Furlan_v4.docx")


# --- UTILITY DI STILE ---
def set_run_font(run, name="Calibri", size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def highlight_xxx(paragraph, text):
    """Aggiunge testo evidenziando 'XXX' in giallo."""
    parts = text.split("XXX")
    for i, part in enumerate(parts):
        if part:
            r = paragraph.add_run(part)
            set_run_font(r)
        if i < len(parts) - 1:
            r = paragraph.add_run("XXX")
            set_run_font(r, bold=True)
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def p(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, space_after=6):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    if "XXX" in text:
        highlight_xxx(para, text)
    else:
        r = para.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return para


def h1(doc, text):
    """Header art. XX bold."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(text)
    set_run_font(r, size=12, bold=True)
    return para


def title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(18)
    r = para.add_run(text)
    set_run_font(r, size=14, bold=True)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    if "XXX" in text:
        highlight_xxx(para, text)
    else:
        r = para.add_run(text)
        set_run_font(r)
    return para


# --- COSTRUZIONE DOCUMENTO ---
doc = Document()

# margini
sec = doc.sections[0]
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)

title(doc, "CONTRATTO DI MANDATO E PRESTAZIONE DI SERVIZI PER LA GESTIONE DI LOCAZIONI TURISTICHE BREVI")

p(doc, 'Il presente contratto di mandato e prestazione di servizi per la gestione di locazioni turistiche brevi (di seguito, il "Contratto"), viene sottoscritto tra:')

# --- MANDANTE 1 ---
p(doc,
  f'{MANDANTE1["nome"]}, nato a {MANDANTE1["nato_luogo"]} il {MANDANTE1["nato_data"]}, '
  f'residente in {MANDANTE1["residenza"]}, C.F. {MANDANTE1["cf"]} '
  '(di seguito, il "Mandante 1");')

p(doc, "e")

# --- MANDANTE 2 ---
p(doc,
  f'{MANDANTE2["nome"]}, nata a {MANDANTE2["nato_luogo"]} il {MANDANTE2["nato_data"]}, '
  f'residente in {MANDANTE2["residenza"]} '
  '(di seguito, il "Mandante 2");')

p(doc,
  'Il Mandante 1 e il Mandante 2, in qualità di comproprietari dell\'Immobile come infra definito, '
  'sono di seguito congiuntamente denominati "i Mandanti" e agiscono in solido tra loro nei confronti '
  'del Mandatario ai sensi degli artt. 1292 e ss. c.c.;')

p(doc, "e")

# --- MANDATARIO ---
p(doc,
  f'{MANDATARIO["nome"]}, nato a {MANDATARIO["nato_luogo"]} il {MANDATARIO["nato_data"]}, '
  f'C.F. {MANDATARIO["cf"]}, domicilio fiscale {MANDATARIO["domicilio"]}, '
  f'titolare della ditta individuale "Angelo Talarico — Property Manager", P.IVA {MANDATARIO["piva"]}, '
  f'codice ATECO {MANDATARIO["ateco"]} (Gestione di immobili per conto terzi), operante sotto il segno '
  f'distintivo / marchio di fatto "{MANDATARIO["marchio"]}" (sito web: {MANDATARIO["sito"]}); recapito PEC '
  f'{MANDATARIO["pec"]} (di seguito, il "Mandatario");')

p(doc, 'i Mandanti e il Mandatario, di seguito collettivamente denominati le "Parti".')

h1(doc, "PREMESSO CHE")

p(doc,
  f'A. I Mandanti sono comproprietari, in regime di comunione pro-indiviso, dell\'immobile sito in '
  f'{IMMOBILE["indirizzo"]}, identificato catastalmente presso il Catasto dei Fabbricati del Comune di '
  f'{IMMOBILE["comune"]} come Sezione urbana {IMMOBILE["sezione_urbana"]}, Foglio {IMMOBILE["foglio"]}, '
  f'Particella {IMMOBILE["particella"]}, Subalterno {IMMOBILE["subalterno"]} — piani {IMMOBILE["piani"]} '
  '(di seguito, l\'"Immobile"), e intendono destinarlo ad attività di locazione per finalità turistiche di '
  'breve durata ai sensi della normativa nazionale (in particolare D.L. 24/04/2017 n. 50 conv. con modif. dalla '
  'L. 21/06/2017 n. 96; L. 30/12/2023 n. 213 — Legge di Bilancio 2024 — art. 1 c. 63, che ha modificato l\'art. 4 '
  'c. 2 del D.L. 50/2017 incrementando l\'aliquota di cedolare secca al 26% dal secondo immobile in poi destinato '
  'a locazione breve nel periodo d\'imposta; D.L. 18/10/2023 n. 145 conv. dalla L. 15/12/2023 n. 191 istitutivo '
  'del CIN; L. 30/12/2020 n. 178 art. 1 c. 595, come modificato dall\'art. 1 c. 17 della L. 30/12/2025 n. 199 '
  '(Legge di Bilancio 2026) con decorrenza dal periodo d\'imposta 2026, in tema di presunzione di attività '
  'imprenditoriale oltre due immobili destinati a locazione breve nel medesimo periodo d\'imposta) e regionale '
  '(L.R. Lombardia 1/10/2015 n. 27 e R.R. Lombardia 5/8/2016 n. 7 in materia di strutture ricettive non '
  'alberghiere).')

p(doc,
  'B. Il Mandatario svolge in via professionale, sotto il marchio di fatto "Host Como", attività di property '
  'management per locazioni turistiche brevi sul territorio del Lago di Como, avvalendosi a tal fine, sotto la '
  'propria esclusiva responsabilità organizzativa e gestionale, di subappaltatori e prestatori di servizi '
  'autonomi secondo quanto previsto all\'art. 2.5.')

p(doc,
  'C. I Mandanti hanno interesse ad affidare al Mandatario lo svolgimento dei Servizi (come infra definiti) alle '
  'condizioni di seguito specificate, conferendo un mandato con rappresentanza limitatamente alle attività di '
  'pubblicazione e gestione operativa dell\'annuncio sui Fornitori di Servizi di Prenotazione Approvati, restando '
  'i Mandanti in ogni caso i diretti e unici locatori degli Ospiti ai fini civilistici e fiscali, ciascuno per la '
  'propria quota.')

p(doc,
  'D. I Mandanti confermano di essere in possesso (ovvero di provvedere all\'ottenimento entro la data di '
  'Pubblicazione dell\'Annuncio) di tutte le autorizzazioni amministrative previste dalla normativa nazionale e '
  'regionale, inclusi senza limitazione: il Codice Identificativo Nazionale (CIN) ai sensi dell\'art. 13-ter del '
  'D.L. 18/10/2023 n. 145 conv. dalla L. 15/12/2023 n. 191 e del D.M. MITUR 6/6/2024, integrato dall\'Avviso '
  'MITUR pubblicato in GU n. 130 P. II del 3/9/2024 (obblighi e sanzioni applicabili dal 1° gennaio 2025 in '
  'forza della proroga disposta con avviso MITUR del 22/10/2024); il Codice Identificativo Regionale Lombardia '
  '(CIR), ove ancora applicabile ai sensi della L.R. 27/2015; la presentazione della SCIA / comunicazione di '
  'inizio attività al Comune competente; l\'iscrizione presso il portale Alloggiati Web della Polizia di Stato '
  'ai sensi dell\'art. 109 TULPS.')

p(doc, "TUTTO CIÒ PREMESSO le Parti convengono e stipulano quanto segue:", bold=True)

# --- ART 1 ---
h1(doc, "1. PREMESSE E ALLEGATI; DEFINIZIONI; PREVALENZA TRA DOCUMENTI")

p(doc, "1.1. Le premesse e i seguenti allegati formano parte integrante e sostanziale del presente Contratto:")
bullet(doc, "Allegato A — Descrizione dei Servizi;")
bullet(doc, "Allegato B — Prerequisiti dell'Immobile e dotazioni obbligatorie di sicurezza;")
bullet(doc, "Allegato C — Tariffe e condizioni economiche;")
bullet(doc, "Allegato D — Trattamento dei dati personali (DPA ex art. 28 GDPR).")

p(doc,
  "1.2. Nel presente Contratto le seguenti espressioni in lettera maiuscola avranno il significato qui di "
  "seguito attribuito, a prescindere dall'impiego al singolare o al plurale:")

DEFS = [
    ('"Account"', 'il profilo creato presso i Fornitori di Servizi di Prenotazione Approvati, intestato e titolato ai Mandanti congiuntamente, sul quale il Mandatario opera in qualità di Co-Host ai sensi dell\'art. 4;'),
    ('"Annuncio"', 'la scheda commerciale dell\'Immobile pubblicata sui Fornitori di Servizi di Prenotazione Approvati, comprensiva di titolo, descrizione, foto, prezzi e calendario;'),
    ('"CIN"', 'il Codice Identificativo Nazionale di cui all\'art. 13-ter del D.L. 18/10/2023 n. 145, convertito con modificazioni dalla L. 15/12/2023 n. 191, e relativi decreti attuativi (in particolare il D.M. MITUR 6/6/2024, integrato dall\'Avviso MITUR pubblicato in GU n. 130 P. II del 3/9/2024, obblighi e sanzioni applicabili dal 1° gennaio 2025 in forza della proroga disposta con avviso MITUR del 22/10/2024);'),
    ('"CIR"', 'il Codice Identificativo Regionale Lombardia previsto dalla L.R. Lombardia 1/10/2015 n. 27 e dai relativi regolamenti attuativi, ove ancora applicabile;'),
    ('"Co-Host"', 'il ruolo formale assegnato al Mandatario sui Fornitori di Servizi di Prenotazione Approvati, con i permessi operativi descritti all\'art. 4;'),
    ('"Fornitore di Servizi di Prenotazione Approvato" (anche "OTA")', 'ciascun soggetto titolare dei diritti sul sito web o piattaforma sulla quale viene pubblicizzato e proposto in locazione l\'Immobile. In via iniziale, le Parti danno atto che saranno utilizzati Airbnb, Booking.com ed Expedia; resta in via esclusiva al Mandatario la facoltà di aggiungere o sostituire ulteriori piattaforme con semplice comunicazione informativa ai Mandanti;'),
    ('"GDPR"', 'il Regolamento (UE) 2016/679 del Parlamento europeo e del Consiglio del 27 aprile 2016;'),
    ('"Host Como"', 'il segno distintivo / marchio di fatto sotto il quale il Mandatario esercita la propria attività professionale nell\'ambito del presente Contratto;'),
    ('"Immobile"', 'ha il significato di cui alla Premessa A;'),
    ('"Mandanti"', 'Andrea Furlan ed Ephanie Sacramento, congiuntamente e in solido, in qualità di comproprietari dell\'Immobile;'),
    ('"Mandatario"', 'Angelo Talarico, titolare della ditta individuale "Angelo Talarico — Property Manager";'),
    ('"Ospiti"', 'i soggetti che effettuano prenotazioni e usufruiscono dell\'Immobile;'),
    ('"Parti"', 'i Mandanti e il Mandatario, collettivamente;'),
    ('"Periodo di Disponibilità"', 'l\'arco temporale in cui l\'Immobile risulta prenotabile su almeno un OTA attivo, pari a non meno di 150 (centocinquanta) giorni in ciascun anno contrattuale. Il computo è effettuato sommando i giorni di calendario in cui l\'Immobile risulta prenotabile su almeno un OTA attivo (computando anche i giorni effettivamente prenotati) e sottraendo unicamente i giorni di indisponibilità tecnica imputabile ai Mandanti; sono parimenti esclusi dal calcolo i giorni bloccati dai Mandanti ai sensi dell\'art. 3.2(e);'),
    ('"Pubblicazione dell\'Annuncio"', 'data di prima pubblicazione attiva dell\'Annuncio su almeno un Fornitore di Servizi di Prenotazione Approvato;'),
    ('"Quota"', 'il corrispettivo complessivo dovuto al Mandatario per l\'esecuzione dei Servizi, secondo quanto descritto all\'art. 7;'),
    ('"Quota di Pernottamento"', 'il corrispettivo lordo di ciascun pernottamento dell\'Immobile in vigenza del Contratto, indipendentemente dal canale di prenotazione (OTA, contatto diretto, contatto privato dei Mandanti o di terzi), comprensivo — per le prenotazioni transitate da un OTA — della commissione applicata dall\'OTA medesima, al netto unicamente delle voci escluse dall\'art. 7.1, lettera a). Per i pernottamenti non fatturati da un OTA, la Quota di Pernottamento è il corrispettivo lordo effettivamente concordato con l\'Ospite, comprensivo di IVA ove applicabile;'),
    ('"Servizi"', 'i servizi indicati all\'art. 2 e dettagliati nell\'Allegato A;'),
    ('"TULPS"', 'il Testo Unico delle Leggi di Pubblica Sicurezza approvato con R.D. 18/06/1931 n. 773 e successive modificazioni, in particolare l\'art. 109 in materia di comunicazione delle generalità degli alloggiati.'),
]
for term, definition in DEFS:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"{term}: ")
    set_run_font(r, bold=True)
    r2 = para.add_run(definition)
    set_run_font(r2)

p(doc,
  "Il riferimento alle fonti normative di ogni grado citate nel presente Contratto si intende effettuato ai "
  "testi vigenti al momento dell'applicazione della disciplina contenuta nelle stesse.")

p(doc,
  "1.3. In caso di conflitto o incongruenza tra i documenti contrattuali, troverà applicazione il seguente "
  "ordine di prevalenza: (1) le previsioni del presente Contratto; (2) gli allegati al presente Contratto.")

# --- ART 2 ---
h1(doc, "2. OGGETTO E CONFERIMENTO DEL MANDATO")

p(doc,
  "2.1. Con il presente Contratto, i Mandanti conferiscono al Mandatario, che accetta, mandato con "
  "rappresentanza limitato alle attività di:")
bullet(doc, "pubblicazione e gestione operativa dell'annuncio dell'Immobile sui Fornitori di Servizi di Prenotazione Approvati;")
bullet(doc, "comunicazione con gli Ospiti, gestione del calendario, dei prezzi e delle prenotazioni;")
bullet(doc, "riscossione, per conto e in nome dei Mandanti quando previsto, dell'imposta di soggiorno secondo le modalità ammesse dal Comune competente.")
p(doc,
  "I Mandanti affidano inoltre al Mandatario, dietro il corrispettivo specificato all'art. 7, l'esecuzione dei "
  "Servizi accessori dettagliati nell'Allegato A.")

p(doc,
  "2.2. Servizi aggiuntivi non compresi nel presente Contratto (a titolo esemplificativo e non esaustivo: "
  "restyling, home staging, shooting fotografico professionale, traduzioni, gestione check-in fuori fascia "
  "oraria standard, deep cleaning, sostituzione biancheria con frequenza superiore allo standard) sono di volta "
  "in volta concordati per iscritto tra le Parti, con specifica del corrispettivo dovuto. Resta in ogni caso "
  "salvo il diritto del Mandatario di rifiutare l'esecuzione di servizi aggiuntivi. Non troverà applicazione "
  "l'art. 1661 c.c.")

p(doc,
  "2.3. In fase di avvio della collaborazione, il Mandatario redigerà a titolo gratuito una lista di acquisti "
  "funzionali ed estetici necessari a rendere l'Immobile adatto alla pubblicazione sul mercato turistico, in "
  "coerenza con i prerequisiti dell'Allegato B. Gli articoli indicati saranno pagati anticipatamente dai "
  "Mandanti e acquistati o installati a cura del Mandatario, senza ulteriore corrispettivo per quest'ultimo.")

p(doc,
  "2.4. Limite di durata della locazione turistica breve (soggiorni oltre 30 giorni). Le Parti danno atto che "
  "il presente mandato, ed il regime sostanziale e fiscale ad esso sotteso (in particolare l'art. 4 del D.L. "
  "50/2017, l'applicabilità della cedolare secca, l'intervento dell'OTA quale sostituto d'imposta e l'esonero "
  "dall'obbligo di registrazione del contratto di locazione), trovano applicazione esclusivamente con "
  "riferimento ai soggiorni di durata non superiore a 30 (trenta) giorni continuativi. I soggiorni di durata "
  "superiore a 30 giorni continuativi non costituiscono locazione turistica breve ai sensi della predetta "
  "normativa, sono esclusi dall'oggetto ordinario del presente mandato e potranno essere gestiti unicamente "
  "previo accordo scritto separato tra le Parti. In tal caso restano integralmente a carico dei Mandanti gli "
  "adempimenti e gli oneri conseguenti al diverso regime applicabile (a titolo esemplificativo: registrazione "
  "del contratto, decadenza dal regime della cedolare secca per locazioni brevi, diverso trattamento fiscale e "
  "contributivo), con esonero del Mandatario da qualsiasi responsabilità per gli adempimenti propri di tale "
  "diverso regime.")

p(doc,
  "2.5. Facoltà di subappalto e di avvalimento di terzi. I Mandanti autorizzano sin d'ora il Mandatario ad "
  "avvalersi, per l'esecuzione dei Servizi, di operatori, dipendenti, collaboratori esterni, professionisti o "
  "subappaltatori autorizzati, espressamente compresi: imprese di pulizia, lavanderie industriali, fotografi "
  "professionisti, tecnici manutentori (idraulici, elettricisti, antennisti), traduttori, fornitori di software "
  "gestionali (PMS, channel manager), nonché prestatori di servizi autonomi di consulenza gestionale strategica "
  "e supporto tecnologico (a titolo esemplificativo, per le attività di studio di mercato, pricing, gestione "
  "tecnica dell'annuncio, PMS/channel manager e reportistica). Tali soggetti operano quali fornitori autonomi "
  "del Mandatario, con piena autonomia organizzativa e fiscale, e non instaurano alcun rapporto contrattuale "
  "diretto con i Mandanti, ai quali il Mandatario resta l'unico soggetto obbligato per l'esecuzione dei "
  "Servizi. La responsabilità verso i Mandanti per l'operato di tali soggetti resta in capo al Mandatario nei "
  "limiti dell'art. 6.")

# --- ART 3 ---
h1(doc, "3. OBBLIGHI DELLE PARTI")

p(doc,
  "3.1. Il Mandatario si impegna ad adempiere alle obbligazioni a proprio carico in modo diligente e secondo "
  "gli standard di mercato. Resta inteso che i Mandanti autorizzano sin d'ora il Mandatario ad avvalersi di "
  "operatori, collaboratori e subappaltatori per l'esecuzione dei Servizi, secondo quanto previsto all'art. 2.5.")

p(doc, "3.2. I Mandanti, in solido tra loro, si impegnano a:")

OBBLIGHI = [
    ("a)", "fornire l'Immobile al Mandatario (i) in buone condizioni d'uso e idoneo alla destinazione turistica, (ii) conforme ai prerequisiti dell'Allegato B (incluse le dotazioni di sicurezza ivi indicate), (iii) libero da qualsivoglia oggetto di valore e/o fragile non espressamente segnalato e custodito;"),
    ("b)", "ottenere e mantenere validi tutti i titoli abilitativi necessari alla locazione turistica breve, in particolare CIN ai sensi dell'art. 13-ter D.L. 145/2023 conv. L. 191/2023 e D.M. MITUR 6/6/2024 (obblighi e sanzioni applicabili dal 1° gennaio 2025 in forza della proroga disposta con avviso MITUR del 22/10/2024), eventuale CIR Lombardia, SCIA al Comune competente, registrazione presso il portale Alloggiati Web della Polizia di Stato ai sensi dell'art. 109 TULPS (con credenziali da rendere disponibili al Mandatario per i relativi adempimenti operativi);"),
    ("c)", "adempiere agli obblighi fiscali gravanti sul locatore, tenuto conto che la presente locazione turistica breve è soggetta al regime fiscale dell'art. 4 del D.L. 50/2017. Le Parti danno atto che, ai sensi dell'art. 1 c. 63 della L. 30/12/2023 n. 213 che ha modificato l'art. 4 c. 2 del D.L. 50/2017, la cedolare secca si applica al 21% per il primo immobile e al 26% dal secondo immobile in poi destinato a locazione breve nel periodo d'imposta; a decorrere dal periodo d'imposta 2026, qualora i Mandanti destinino alla locazione breve più di due immobili nel corso del medesimo periodo d'imposta, l'attività si presume esercitata in forma imprenditoriale ai sensi dell'art. 1 c. 595 L. 178/2020, come modificato dall'art. 1 c. 17 della L. 30/12/2025 n. 199 (Legge di Bilancio 2026), con conseguente obbligo dal terzo immobile di apertura della partita IVA, iscrizione al Registro delle Imprese mediante Comunicazione Unica, presentazione di SCIA, assoggettamento ad IVA dei corrispettivi e perdita del regime della cedolare secca; IVA in regime ordinario o forfetario ove dovuta; IMU e tributi locali ove dovuti;"),
    ("d)", "autorizzare il Mandatario, secondo quanto previsto all'art. 4, ad operare in qualità di Co-Host sull'Account, conservando i Mandanti la titolarità formale dell'Account, dell'annuncio e di ogni contenuto multimediale ad esso correlato;"),
    ("e)", "rendere disponibile l'Immobile per il Periodo di Disponibilità (almeno 150 giorni/anno). I Mandanti possono riservare l'Immobile per sé o per persone da loro indicate previo preavviso al Mandatario di almeno 15 (quindici) giorni, salvo che il periodo non sia già oggetto di prenotazione confermata di un Ospite;"),
    ("f)", "fornire tempestivamente al Mandatario informazioni complete, accurate e puntuali sull'Immobile e su qualsiasi ulteriore aspetto che possa risultare necessario per la corretta esecuzione dei Servizi;"),
    ("g)", "consegnare al Mandatario almeno n. 3 (tre) mazzi completi di chiavi dell'Immobile (o credenziali equivalenti per serrature smart);"),
    ("h)", "sostenere tutti i costi necessari per: (i) le operazioni di manutenzione ordinaria e straordinaria dell'Immobile; (ii) l'adeguamento ai prerequisiti dell'Allegato B (inclusi arredi e dotazioni di sicurezza); (iii) i tributi locali e nazionali gravanti sull'Immobile e sull'attività locativa; (iv) le forniture (elettricità, gas, acqua, internet, gestione rifiuti); (v) i danni causati dagli Ospiti agli arredi e alle strutture nei limiti non coperti dalle eventuali polizze o depositi cauzionali; (vi) la normale usura dei beni;"),
    ("i)", "stipulare e mantenere in essere, per tutta la durata del Contratto, un'unica polizza assicurativa intestata congiuntamente ad entrambi i Mandanti a copertura (i) della responsabilità civile verso terzi del proprietario dell'Immobile, comprensiva dei danni alla persona di Ospiti e terzi, e (ii) dei danni all'immobile e ai contenuti causati da locazioni brevi (polizza \"casa vacanze\" o estensione locazione breve della polizza globale fabbricato), con massimale per sinistro non inferiore a Euro 1.000.000,00 (un milione/00). I Mandanti richiedono al proprio assicuratore l'inserimento del Mandatario quale \"assicurato aggiuntivo\" (additional insured) ovvero, ove la specifica forma contrattuale di polizza non lo consenta, quale \"vincolatario\" o beneficiario, in relazione alle attività di gestione svolte presso l'Immobile, ed esibiscono al Mandatario copia integrale del testo di polizza prima della Pubblicazione dell'Annuncio per la verifica dell'effettiva inclusione e della portata oggettiva delle garanzie. Ove l'assicuratore non conceda alcuna delle predette estensioni, i Mandanti ne danno tempestiva comunicazione scritta al Mandatario, restando ferma l'efficacia del sistema integrato di copertura di cui all'art. 6.4 quale unica protezione disponibile in tal caso. Copia della polizza in corso di validità è esibita al Mandatario su richiesta scritta entro 15 (quindici) giorni dalla richiesta medesima;"),
    ("j)", "rispettare tempestivamente le ulteriori obbligazioni poste a proprio carico negli Allegati al presente Contratto."),
]
for lett, testo in OBBLIGHI:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"{lett} ")
    set_run_font(r, bold=True)
    r2 = para.add_run(testo)
    set_run_font(r2)

p(doc,
  "I Mandanti riconoscono espressamente che il mancato adempimento, anche solo di una delle obbligazioni di cui "
  "al presente art. 3.2, può causare ritardi o impedire la corretta esecuzione dei Servizi. In tal caso, il "
  "Mandatario non potrà essere ritenuto responsabile per (i) il non corretto adempimento dei propri obblighi e "
  "(ii) eventuali richieste di risarcimento avanzate dagli Ospiti o da terzi in genere.")

p(doc,
  "3.3. Una volta confermata la prenotazione di ciascun Ospite, i Mandanti si impegnano a non cancellarla e/o "
  "modificarla, fatto salvo il caso in cui le condizioni applicate dal Fornitore di Servizi di Prenotazione "
  "Approvato lo permettano. In caso di modifica e/o cancellazione non consentita, i Mandanti in solido saranno "
  "tenuti a pagare la Quota al Mandatario come se la prenotazione fosse stata regolarmente eseguita, fatto salvo "
  "il diritto del Mandatario di risolvere il presente Contratto ai sensi dell'art. 1456 c.c.")

p(doc, "3.4. Il Mandatario si impegna a:")
OBBLIGHI_M = [
    ("a)", "gestire in autonomia tutte le richieste e problematiche degli Ospiti (a titolo esemplificativo: richieste di assistenza, malfunzionamenti del wifi, smarrimento chiavi, comunicazioni in più lingue);"),
    ("b)", "effettuare per conto dei Mandanti la registrazione degli Ospiti sul portale Alloggiati Web della Polizia di Stato ai sensi dell'art. 109 TULPS, entro 24 (ventiquattro) ore dall'arrivo dell'Ospite ovvero, qualora il soggiorno abbia durata inferiore alle 24 ore, entro le 6 (sei) ore successive all'arrivo, utilizzando le credenziali fornite dai Mandanti;"),
    ("c)", "effettuare le comunicazioni mensili ISTAT per conto dei Mandanti secondo le modalità previste dalla Regione Lombardia e dal Comune competente;"),
    ("d)", "riscuotere l'imposta di soggiorno dagli Ospiti secondo le modalità ammesse dal Comune competente, con preferenza per la riscossione tramite OTA ove disponibile (art. 7.4), tracciandola separatamente e versandola periodicamente ai Mandanti affinché provvedano al riversamento al Comune entro i termini di legge. Il Mandatario fornisce ai Mandanti con cadenza trimestrale il riepilogo delle imposte di soggiorno raccolte;"),
    ("e)", "redigere e inviare ai Mandanti un resoconto mensile dettagliato delle prenotazioni ricevute, dei pagamenti incassati dai Mandanti tramite OTA, dei pagamenti dovuti al Mandatario e ai suoi operatori/subappaltatori, nonché di eventuali compensazioni;"),
    ("f)", "comunicare tempestivamente ai Mandanti eventuali danni all'Immobile riscontrati dopo la partenza degli Ospiti, attivando ove possibile i rimborsi previsti dalle piattaforme (es. AirCover di Airbnb)."),
]
for lett, testo in OBBLIGHI_M:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"{lett} ")
    set_run_font(r, bold=True)
    r2 = para.add_run(testo)
    set_run_font(r2)

p(doc,
  "3.5. Per tutta la durata del Contratto, i Mandanti si obbligano a non affidare la gestione dell'Immobile "
  "per finalità di locazione turistica breve a soggetti diversi dal Mandatario, né a operare in proprio sulla "
  "stessa piattaforma utilizzata dal Mandatario per la pubblicazione dell'Annuncio, salvo diverso accordo "
  "scritto tra le Parti. Restano ferme le previsioni dell'art. 7.1-bis in materia di prenotazioni dirette: i "
  "Mandanti non sono obbligati a rinunciare ai propri contatti privati, ma ogni pernottamento dell'Immobile a "
  "titolo oneroso, da qualunque canale provenga, è soggetto alla Quota. Resta altresì ferma la facoltà dei "
  "Mandanti di creare e promuovere un proprio sito web dell'Immobile ai sensi dell'art. 7.1-ter, non "
  "costituendo tale attività violazione del presente vincolo di esclusiva. La violazione del presente vincolo "
  "di esclusiva costituisce grave inadempimento e legittima il Mandatario alla risoluzione del Contratto ex "
  "art. 1456 c.c., fatto salvo il risarcimento del danno.")

p(doc,
  "3.6. Conseguenze del mancato ottenimento del CIN. I Mandanti si impegnano ad ottenere e mantenere valido il "
  "Codice Identificativo Nazionale (CIN) di cui all'art. 13-ter D.L. 145/2023 conv. L. 191/2023 entro la data "
  "di Pubblicazione dell'Annuncio. In mancanza, l'Annuncio non potrà essere attivato dal Mandatario; ove sia "
  "stato attivato in via provvisoria su un canale OTA che lo consenta, lo stesso viene automaticamente sospeso "
  "sino all'ottenimento del CIN, senza che da ciò derivi alcuna responsabilità in capo al Mandatario. Le Parti "
  "danno reciprocamente atto che la pubblicazione di annuncio privo di CIN espone i Mandanti, in caso di "
  "assenza del CIN o di mancata richiesta dello stesso, a sanzione amministrativa pecuniaria da Euro 800,00 a "
  "Euro 8.000,00 ai sensi dell'art. 13-ter c. 9, primo periodo, D.L. 145/2023 conv. L. 191/2023, determinata "
  "in relazione alle dimensioni dell'immobile (resta ferma, in caso di mancata esposizione o indicazione del "
  "CIN nell'annuncio, la diversa sanzione da Euro 500,00 a Euro 5.000,00 per ciascuna unità immobiliare, con "
  "immediata rimozione dell'annuncio irregolare), oltre al rischio di rimozione dell'Annuncio da parte "
  "dell'OTA. Qualora i Mandanti non ottengano il CIN entro 90 (novanta) giorni dalla data di sottoscrizione "
  "del presente Contratto, ciascuna Parte può risolvere il Contratto ai sensi dell'art. 1456 c.c. senza "
  "preavviso e senza applicazione della penale di cui all'art. 8.2-bis.")

p(doc,
  "3.7. Procedura di gestione dei danni non coperti da polizza o piattaforma. In caso di danno all'Immobile, "
  "agli arredi o ai contenuti causato dagli Ospiti e non coperto integralmente da AirCover (o programmi di "
  "protezione equivalenti dei Fornitori di Servizi di Prenotazione Approvati) né dalla polizza assicurativa dei "
  "Mandanti di cui all'art. 3.2(i), il Mandatario procede come segue: (i) entro 48 (quarantotto) ore dalla "
  "scoperta del danno, segnala l'evento ai Mandanti con documentazione fotografica e una stima dei costi di "
  "ripristino; (ii) i Mandanti, entro 5 (cinque) giorni dalla segnalazione, comunicano per iscritto se "
  "intendono attivare la propria polizza, trattenere eventuali depositi cauzionali / cauzioni d'incasso o "
  "sostenere direttamente il costo; (iii) decorso il termine senza riscontro, il Mandatario procede con la "
  "modalità ritenuta meno onerosa per i Mandanti, addebitando i relativi costi nel resoconto mensile.")

p(doc,
  "3.8. Checklist di sicurezza e rilievo dello stato dell'Immobile all'onboarding. Prima della Pubblicazione "
  "dell'Annuncio e comunque non oltre 15 (quindici) giorni dalla sottoscrizione del Contratto, i Mandanti (o "
  "almeno uno di essi, in rappresentanza dell'altro) e il Mandatario procedono congiuntamente al rilievo "
  "formale dello stato dell'Immobile mediante: (i) reportage fotografico documentale con datazione automatica "
  "di tutti i locali, impianti tecnici, dotazioni di sicurezza e arredi; (ii) compilazione e sottoscrizione, "
  "da parte dei Mandanti (o del Mandante presente in rappresentanza dell'altro), di una checklist di sicurezza "
  "congruente con i prerequisiti dell'Allegato B, comprensiva tra l'altro della verifica e attestazione dei "
  "Mandanti in ordine a: conformità dell'impianto elettrico (con riferimento alla relativa dichiarazione di "
  "rispondenza ex D.M. 37/2008), conformità dell'impianto termico, presenza e funzionalità di rilevatori di "
  "fumo e di monossido di carbonio (CO), presenza ed efficienza di estintori, presenza di interruttore "
  "differenziale (\"salvavita\") e di idoneo impianto di messa a terra, agibilità delle vie di fuga, idoneità "
  "delle dotazioni minime ex Allegato B. La checklist sottoscritta e il reportage fotografico costituiscono "
  "parte integrante del Contratto e fanno prova, salvo querela di falso, dello stato dell'Immobile alla data "
  "di attivazione ai fini del riparto di responsabilità di cui agli artt. 3.2 lettera a), 5.2 e 6.1 lettera a). "
  "I Mandanti sono tenuti a comunicare per iscritto al Mandatario, senza ritardo, ogni variazione successiva "
  "dello stato dei sistemi di sicurezza o degli impianti, nonché l'esito delle verifiche periodiche "
  "obbligatorie ai sensi della normativa applicabile.")

# --- ART 4 ---
h1(doc, "4. ACCOUNT, RUOLI E TITOLARITÀ DEI CONTENUTI")

p(doc,
  "4.1. Il Mandatario, su delega dei Mandanti, provvede alla creazione dell'Account presso i Fornitori di "
  "Servizi di Prenotazione Approvati. In via iniziale, sono attivati gli Account presso Airbnb, Booking.com ed "
  "Expedia. L'Account è intestato ai Mandanti congiuntamente, che ne sono titolari formali e \"Host\" verso i "
  "terzi. Tutte le credenziali (username, email principale, password di recupero) sono comunicate ai Mandanti "
  "e loro appartengono.")

p(doc,
  "4.2. Il Mandatario ha piena discrezionalità nell'attivare ulteriori Fornitori di Servizi di Prenotazione "
  "Approvati (a titolo esemplificativo: Vrbo, HomeAway, TripAdvisor Rentals, agenzie locali, sito proprietario "
  "hostcomo.com) o nel sospendere/disattivare canali che si rivelino inefficaci, con semplice comunicazione "
  "informativa ai Mandanti. L'attivazione di nuovi canali non comporta alcun costo aggiuntivo a carico dei "
  "Mandanti, fatte salve le commissioni applicate direttamente dalla piattaforma sulle prenotazioni, coerenti "
  "con quanto previsto all'art. 7. Le indicazioni dei Mandanti ai sensi dell'art. 5.1, lett. a), hanno valore "
  "consultivo e non vincolante; il Mandatario le valuta in buona fede ma decide autonomamente nell'interesse "
  "della performance dell'Immobile.")

p(doc,
  "4.3. I Mandanti conferiscono al Mandatario il ruolo di \"Co-Host\" (o ruolo analogo previsto dalla "
  "piattaforma) con permessi operativi pieni per: pubblicare e modificare l'annuncio, gestire calendario e "
  "prezzi, comunicare con gli Ospiti, gestire le prenotazioni e gli incassi (che restano accreditati al conto "
  "designato dai Mandanti), rispondere alle recensioni.")

p(doc,
  "4.4. I pagamenti effettuati dagli Ospiti affluiscono direttamente sul conto bancario designato dai Mandanti, "
  "comunicato all'OTA al momento della configurazione dell'Account. Il Mandatario non riceve né detiene somme "
  "spettanti ai Mandanti in relazione alle Quote di Pernottamento e, conseguentemente, non opera in qualità di "
  "sostituto d'imposta ai sensi dell'art. 4, c. 5-bis del D.L. 50/2017 conv. L. 96/2017. I Mandanti restano gli "
  "unici responsabili della corretta dichiarazione e versamento delle imposte sui canoni di locazione.")

p(doc,
  "4.5. L'annuncio, le foto (anche se realizzate dal Mandatario o da fotografi da lui incaricati), le "
  "descrizioni, le recensioni accumulate e ogni altro contenuto multimediale o testuale pubblicato "
  "sull'Account rimangono di titolarità congiunta dei Mandanti. Alla cessazione del Contratto, tali contenuti "
  "restano nella piena disponibilità dei Mandanti senza necessità di ulteriori atti di trasferimento. Il "
  "Mandatario rinuncia espressamente a ogni rivendicazione di diritti d'autore o connessi su tali contenuti, "
  "fermo restando il diritto di portafoglio professionale (uso interno e dimostrativo non commerciale).")

p(doc,
  "4.6. Alla cessazione del Contratto per qualunque causa, i Mandanti revocheranno tempestivamente i permessi "
  "di Co-Host del Mandatario e questi cesserà ogni operatività sull'Account.")

# --- ART 5 ---
h1(doc, "5. DICHIARAZIONI DEI MANDANTI E MANLEVA")

p(doc, "5.1. I Mandanti riconoscono e dichiarano:")
DICH = [
    ("a)", "di essere informati e accettare che il Mandatario agisce in qualità di gestore-mandatario in forza dei poteri qui conferiti, e non assume alcun obbligo di consultare i Mandanti prima di accettare una prenotazione o fissarne il prezzo, fatta salva la facoltà dei Mandanti di esprimere indicazioni o suggerimenti;"),
    ("b)", "di essere informati e accettare che il Mandatario non fornisce consulenza in materia di investimento immobiliare né di pianificazione fiscale, e che qualsiasi decisione di tale natura è ascrivibile esclusivamente ai Mandanti (ciascuno per la propria quota);"),
    ("c)", "di aver verificato l'idoneità dell'Immobile alla destinazione turistica e di aver ottenuto, ovvero di impegnarsi a ottenere prima della Pubblicazione dell'Annuncio, tutti i titoli amministrativi necessari."),
]
for lett, testo in DICH:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"{lett} ")
    set_run_font(r, bold=True)
    r2 = para.add_run(testo)
    set_run_font(r2)

p(doc,
  "5.2. I Mandanti si impegnano, in solido tra loro, a manlevare e tenere indenne il Mandatario da ogni "
  "pretesa, costo, responsabilità, danno, perdita e spesa (incluse le ragionevoli spese legali) derivanti da "
  "fatti riconducibili alla sfera di controllo dei Mandanti medesimi, e in particolare da: (i) violazione dei "
  "propri obblighi ai sensi del presente Contratto, ivi compresi gli obblighi di cui all'art. 3.2; (ii) "
  "carenze nei titoli amministrativi (CIN, CIR, SCIA, conformità impianti, dotazioni di sicurezza di cui "
  "all'Allegato B); (iii) difetti strutturali, di sicurezza o di conformità normativa dell'Immobile; (iv) "
  "richieste di risarcimento degli Ospiti correlate a danni cagionati dagli Ospiti stessi. La manleva di cui "
  "al presente art. 5.2 non opera in caso di dolo o colpa grave del Mandatario o dei suoi "
  "operatori/subappaltatori, di violazione di obblighi inderogabili di legge a carico del Mandatario, né in "
  "caso di danni alla persona ai sensi dell'art. 1229 c.c., nei limiti di quanto previsto all'art. 6.")

p(doc,
  "5.3. Qualificazione dei Mandanti (consumatore / professionista). Ciascun Mandante dichiara di sottoscrivere "
  "il presente Contratto:")
para = doc.add_paragraph()
r = para.add_run("☒ ")
set_run_font(r, bold=True)
r2 = para.add_run(
    "quale consumatore ai sensi dell'art. 3 del D.Lgs. 206/2005 (Codice del Consumo), agendo per scopi "
    "estranei all'attività imprenditoriale, commerciale, artigianale o professionale eventualmente svolta;")
set_run_font(r2)
para = doc.add_paragraph()
r = para.add_run("☐ ")
set_run_font(r, bold=True)
r2 = para.add_run("nell'esercizio della propria attività imprenditoriale o professionale.")
set_run_font(r2)
p(doc, "(barrare la casella pertinente).")

p(doc,
  "Le Parti danno reciprocamente atto che, qualora i Mandanti agiscano in qualità di consumatori, le clausole "
  "eventualmente vessatorie restano soggette alla disciplina inderogabile di cui agli artt. 33-36 del Codice "
  "del Consumo, che prevale sull'accettazione specifica ex artt. 1341-1342 c.c.; resta fermo il foro "
  "inderogabile del consumatore di cui all'art. 17.3.")

# --- ART 6 ---
h1(doc, "6. LIMITAZIONI DI RESPONSABILITÀ DEL MANDATARIO")

p(doc,
  "6.1. Nella misura massima prevista dalla legge applicabile, i Mandanti espressamente accettano e "
  "riconoscono che il Mandatario non potrà essere ritenuto responsabile per:")
LIM = [
    ("a)", "la sicurezza, le condizioni strutturali, la conformità normativa e la manutenzione straordinaria dell'Immobile, che restano integralmente in capo ai Mandanti;"),
    ("b)", "eventuali rapporti diretti intercorsi tra i Mandanti e i Fornitori di Servizi di Prenotazione Approvati al di fuori del perimetro delle operazioni gestite dal Mandatario in qualità di Co-Host;"),
    ("c)", "comportamenti, danni, furti o atti illeciti compiuti dagli Ospiti, salvo i casi di colpa grave del Mandatario nella selezione e accettazione della prenotazione;"),
    ("d)", "la fornitura di servizi da parte di imprese terze (impresa di pulizie, lavanderia, tecnici manutentori) quando il Mandatario abbia operato con la diligenza richiesta nella loro selezione e coordinamento;"),
    ("e)", "violazione, ritardo nell'esecuzione o mancata esecuzione dei propri obblighi dovuti a cause che esulino dal proprio controllo (a titolo esemplificativo: indisponibilità dell'Immobile per causa imputabile ai Mandanti, overbooking della piattaforma, condizioni climatiche avverse, interruzione di utenze, disastri naturali, scioperi, atti di terrorismo, guerra, sommosse, pandemie, blocchi normativi)."),
]
for lett, testo in LIM:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(f"{lett} ")
    set_run_font(r, bold=True)
    r2 = para.add_run(testo)
    set_run_font(r2)

p(doc,
  "6.2. Il Mandatario è economicamente responsabile per i danni causati direttamente da lui o dai propri "
  "operatori/subappaltatori all'Immobile, agli arredi e ai contenuti, nei limiti dell'ordinaria diligenza "
  "professionale. Restano esclusi i danni derivanti dalla normale usura dei beni e dall'utilizzo degli Ospiti.")

p(doc,
  "6.3. Fermo quanto previsto agli artt. 6.1 e 6.2, la responsabilità complessiva del Mandatario in "
  "riferimento al presente Contratto è limitata all'importo più elevato tra (i) il costo sostenuto dai "
  "Mandanti per l'ottenimento di servizi sostitutivi, (ii) la Quota media percepita dal Mandatario nei 12 "
  "(dodici) mesi precedenti la data dell'evento dannoso e (iii) un floor minimo di Euro 5.000,00 "
  "(cinquemila/00). Il presente limite non opera in caso di dolo o colpa grave del Mandatario o dei suoi "
  "operatori/subappaltatori, né in caso di danni alla persona o di violazione di norme inderogabili di legge, "
  "ai sensi dell'art. 1229 c.c. Considerato che entrambi i Mandanti sono qualificati come consumatori ai sensi "
  "dell'art. 5.3, il cap di responsabilità di cui al presente articolo non opera e la responsabilità del "
  "Mandatario è regolata dalle norme di legge applicabili, restando ferme le coperture assicurative di cui "
  "agli artt. 3.2 lettera i) e 6.4.")

p(doc,
  "6.4. Polizze assicurative degli operatori e schermo dei rischi alla persona. Tenuto conto che i Servizi di "
  "esecuzione materiale (pulizie, lavanderia, manutenzioni, accoglienza, fotografia, prestazioni autonome di "
  "cui all'art. 2.5) sono resi ai Mandanti per il tramite di operatori e fornitori indipendenti selezionati "
  "dal Mandatario, le Parti convengono che ciascuno di tali operatori è tenuto a mantenere in essere, in "
  "proprio, polizza assicurativa di responsabilità civile professionale adeguata all'attività svolta e idonea "
  "a coprire i danni a persone e cose conseguenti all'esercizio della propria prestazione. Il Mandatario "
  "verifica la sussistenza di tale copertura in fase di selezione e ne esige il rinnovo periodico mediante "
  "(i) richiesta e archiviazione di copia della polizza in corso di validità del fornitore o di certificato "
  "assicurativo equivalente, (ii) acquisizione di dichiarazione sostitutiva di atto di notorietà del fornitore "
  "attestante l'assenza di sospensioni, disdette o esclusioni rilevanti, (iii) re-verifica almeno annuale o in "
  "occasione di ogni nuovo affidamento di prestazione di valore non occasionale. Qualora un operatore o "
  "fornitore non risulti coperto da polizza idonea, il Mandatario è tenuto a sostituirlo entro 30 (trenta) "
  "giorni con altro fornitore conforme, salvo che la prestazione sia di carattere occasionale, di valore "
  "esiguo e tale da non comportare rischi significativi per persone o cose, ipotesi nella quale il Mandatario "
  "ne dà comunque comunicazione preventiva ai Mandanti, restando esonerato — fermo il dovere di diligenza "
  "nella scelta e nel coordinamento ai sensi dell'art. 1228 c.c. — da responsabilità diretta per fatti dei "
  "medesimi operatori che esorbitino dal proprio controllo organizzativo. Il Mandatario, salva diversa "
  "autonoma valutazione, non è obbligato dal presente Contratto a stipulare a propria volta polizza "
  "assicurativa di responsabilità civile professionale, restando il rischio dei danni alla persona coperto dal "
  "sistema integrato di (a) polizza dei Mandanti di cui all'art. 3.2 lettera i), (b) polizze proprie degli "
  "operatori indipendenti come sopra, (c) programmi di protezione dei Mandanti offerti dai Fornitori di "
  "Servizi di Prenotazione Approvati (a titolo esemplificativo, AirCover di Airbnb). Le Parti riconoscono "
  "espressamente che il presente sistema integrato di copertura distribuisce ma non elimina il rischio dei "
  "danni alla persona di Ospiti o di terzi, restando ferma — fermo l'art. 1229 c.c. che esclude qualunque "
  "limitazione della responsabilità per dolo e colpa grave — la facoltà del Mandatario di stipulare, in "
  "qualunque momento e in funzione della propria valutazione di esposizione al rischio, una polizza propria "
  "di responsabilità civile professionale ad integrazione delle coperture sopra elencate.")

p(doc,
  "6.5. Registro di diligenza operativa. Il Mandatario tiene per l'Immobile un registro digitale (sistema di "
  "ticketing o equivalente strumento tracciato e immodificabile) nel quale annota tempestivamente: (i) le "
  "segnalazioni e i reclami degli Ospiti, (ii) gli incidenti, gli infortuni o i danni a persone o cose di "
  "qualunque entità riscontrati, (iii) le richieste di intervento ai fornitori e i tempi di esecuzione, "
  "(iv) le verifiche periodiche delle dotazioni di sicurezza di cui all'art. 3.8, (v) le comunicazioni "
  "rilevanti scambiate con i Mandanti in ordine allo stato dell'Immobile e dei suoi impianti. Il registro, "
  "conservato per l'intera durata del Contratto e per i 10 (dieci) anni successivi alla cessazione del "
  "medesimo ai fini probatori, è messo a disposizione dei Mandanti, dei loro consulenti, degli assicuratori e "
  "dell'autorità competente su richiesta motivata. Le annotazioni del registro, formate nel rispetto della "
  "diligenza professionale richiesta al property manager, costituiscono elemento di prova della tempestività e "
  "dell'adeguatezza delle azioni intraprese dal Mandatario in relazione a ciascun evento, e ai sensi degli "
  "artt. 2710 e 2712 c.c. assumono rilievo probatorio nei rapporti tra le Parti.")

# --- ART 7 ---
h1(doc, "7. PAGAMENTI E FLUSSI ECONOMICI")

p(doc,
  "7.1. I Mandanti riconoscono al Mandatario, a titolo di corrispettivo complessivo per i Servizi oggetto del "
  "presente Contratto, una commissione (la \"Quota\") pari al 10% (dieci per cento) della Quota di "
  "Pernottamento, oltre IVA dove applicabile in base al regime fiscale del Mandatario.")

p(doc,
  "La Quota è calcolata sulla Quota di Pernottamento come definita all'art. 1, ovvero sul corrispettivo lordo "
  "di ciascun pernottamento dell'Immobile (inclusa la commissione applicata dall'OTA medesima, ove la "
  "prenotazione transiti da un OTA, che resta a carico dei Mandanti secondo le condizioni del singolo canale; "
  "ovvero, per le prenotazioni dirette, sul corrispettivo lordo concordato con l'Ospite), al netto unicamente "
  "delle voci escluse alla successiva lettera a).")

p(doc,
  "a) Voci escluse dalla base di calcolo della Quota. Sono in ogni caso escluse dalla base di calcolo della "
  "Quota: (i) la cleaning fee (corrispettivo di pulizia incassato dall'OTA sull'Ospite per essere riversato al "
  "fornitore del servizio o ai Mandanti); (ii) l'imposta di soggiorno comunale; (iii) eventuali cauzioni e "
  "depositi cauzionali; (iv) i Servizi aggiuntivi di cui all'art. 2.2 ove specificatamente fatturati a parte; "
  "(v) ogni altro corrispettivo di terzi (manutenzioni, ricambi, materiali di consumo) fatturato direttamente "
  "ai Mandanti.")

p(doc,
  "b) Soglia di disponibilità (150 giorni). La percentuale base del 10% è valida a condizione che l'Immobile "
  "sia reso disponibile per la locazione turistica per non meno di 150 (centocinquanta) giorni nell'arco di "
  "ciascun anno contrattuale. Qualora l'Immobile sia reso disponibile per un periodo inferiore a 150 "
  "(centocinquanta) giorni nell'arco di ciascun anno contrattuale, la Quota sulle prenotazioni effettivamente "
  "realizzate è automaticamente incrementata di 2,5 (due virgola cinque) punti percentuali, passando dal 10% "
  "al 12,5% (dodici virgola cinque per cento) oltre IVA dove applicabile, a decorrere dal momento in cui la "
  "soglia viene mancata. La disponibilità dell'Immobile viene calcolata sulla base dei 12 (dodici) mesi "
  "decorrenti dalla data di Pubblicazione dell'Annuncio; la disponibilità non deve necessariamente essere "
  "continuativa.")

p(doc,
  "Eventuali spese e/o costi aggiuntivi sostenuti dal Mandatario su richiesta o per conto dei Mandanti sono "
  "addebitati separatamente nel resoconto mensile.")

p(doc,
  "7.1-bis. Prenotazioni dirette e uso personale dei Mandanti. I Mandanti possono accogliere prenotazioni "
  "provenienti da contatti diretti (repeat guest, conoscenti, contatti privati). Per tali prenotazioni la "
  "Quota di cui all'art. 7.1 si applica esclusivamente nei casi in cui i Servizi siano effettivamente resi "
  "dal Mandatario — in particolare: registrazione Alloggiati Web/CIN, gestione di check-in e check-out, "
  "blocco e gestione della disponibilità sul channel-manager, comunicazione con l'Ospite tramite i canali "
  "del Mandatario, raccolta e versamento dell'imposta di soggiorno, reportistica — ed è calcolata sul "
  "corrispettivo lordo concordato con l'Ospite. La Quota non si applica alle prenotazioni dirette gestite "
  "integralmente in autonomia dai Mandanti, a condizione che i Mandanti (a) comunichino per iscritto al "
  "Mandatario il blocco di disponibilità con preavviso non inferiore a 15 (quindici) giorni, ai fini della "
  "cancellazione su channel-manager, (b) provvedano in proprio agli adempimenti Alloggiati Web/CIN e "
  "all'eventuale dichiarazione e versamento dell'imposta di soggiorno presso il Comune competente, (c) "
  "tengano indenne il Mandatario da ogni responsabilità per la gestione dell'Ospite. Sono in ogni caso esclusi "
  "dall'applicazione della Quota (i) i soggiorni gratuiti personali dei Mandanti medesimi presso l'Immobile e "
  "(ii) i soggiorni gratuiti dei familiari dei Mandanti, intendendosi per tali il coniuge, la parte "
  "dell'unione civile costituita ai sensi della L. 20/5/2016 n. 76, il convivente more uxorio o parte di "
  "convivenza di fatto registrata ai sensi della medesima L. 76/2016, i figli, i genitori e i fratelli o le "
  "sorelle di ciascun Mandante, purché complessivamente (a) comunicati per iscritto al Mandatario con "
  "preavviso non inferiore a 15 (quindici) giorni prima del check-in, ai fini della cancellazione della "
  "disponibilità sul channel-manager e degli adempimenti Alloggiati Web/CIN, (b) per periodi complessivi non "
  "superiori a 30 (trenta) giorni di calendario per anno solare cumulando uso personale dei Mandanti e "
  "soggiorni dei familiari, (c) senza corresponsione di alcun corrispettivo diretto o indiretto. L'omessa "
  "comunicazione preventiva di una prenotazione diretta gestita in autonomia, o di un soggiorno escluso, "
  "comporta l'applicazione della Quota sul corrispettivo medio degli ultimi 12 mesi per pernottamento "
  "equivalente, fatti salvi gli ulteriori rimedi contrattuali in caso di violazione reiterata.")

p(doc,
  "7.1-ter. Sito web dei Mandanti e prenotazioni dirette via canale proprio. I Mandanti sono espressamente "
  "autorizzati a creare, mantenere e promuovere un proprio sito web dedicato all'Immobile per finalità di "
  "prenotazione diretta, anche con incassi gestiti tramite propri prestatori di servizi di pagamento (PSP), e "
  "a gestire integralmente in autonomia il rapporto con l'Ospite (comunicazioni, check-in e check-out, "
  "eventuali adempimenti Alloggiati Web/CIN e imposta di soggiorno per la singola prenotazione, ove non resi "
  "dal Mandatario). I Mandanti sono tenuti a (i) integrare il proprio sito con il channel-manager del "
  "Mandatario ai fini di sincronizzazione del calendario e prevenzione di overbooking, (ii) comunicare "
  "tempestivamente al Mandatario ciascuna prenotazione confermata con dati completi dell'Ospite e del "
  "soggiorno. Per ciascuna prenotazione ricevuta tramite il sito proprio dei Mandanti, anche ove i Mandanti "
  "gestiscano integralmente in autonomia il rapporto con l'Ospite ai sensi del presente articolo, resta "
  "dovuta al Mandatario una Quota pari al 10% (dieci per cento), calcolata — in deroga alla definizione di "
  "\"Quota di Pernottamento\" del Glossario e in deroga all'art. 7.1-bis, secondo periodo — sul corrispettivo "
  "netto effettivamente incassato dai Mandanti, intendendosi per netto il corrispettivo lordo concordato con "
  "l'Ospite al netto delle sole commissioni di pagamento applicate dai PSP utilizzati dai Mandanti (es. "
  "Stripe, PayPal, gateway carte), restando esclusi dalla detrazione costi di hosting del sito, attività di "
  "marketing, SEO o altri costi commerciali dei Mandanti. La Quota così determinata remunera in particolare "
  "l'uso del channel-manager del Mandatario per la sincronizzazione calendario, la prevenzione "
  "dell'overbooking e la conseguente attività di reportistica e amministrazione correlata.")

p(doc,
  "7.2. Il Mandatario emette ai Mandanti un'unica fattura mensile per la Quota complessiva maturata nel mese "
  "di riferimento, indicando in dettaglio le prestazioni rese. La fattura è intestata al Mandante 1 (Andrea "
  "Furlan), fatta salva diversa istruzione scritta congiunta dei Mandanti. I rapporti economici tra il "
  "Mandatario e i propri operatori, fornitori o subappaltatori (ivi inclusi i prestatori di servizi autonomi "
  "di cui all'art. 2.5) restano interni al Mandatario e non riguardano i Mandanti, che restano obbligati "
  "unicamente verso il Mandatario in solido tra loro.")

p(doc, "7.3. Le Parti concordano espressamente che:")

p(doc,
  "a) i pagamenti delle Quote di Pernottamento effettuati dagli Ospiti affluiscono direttamente sul conto "
  "bancario designato dai Mandanti tramite gli strumenti predisposti dai Fornitori di Servizi di Prenotazione "
  "Approvati. Le Parti danno reciprocamente atto, a titolo informativo, che ai sensi dell'art. 4 c. 5-bis del "
  "D.L. 50/2017 i Fornitori di Servizi di Prenotazione Approvati che intervengono nell'incasso o nel pagamento "
  "dei canoni operano in qualità di sostituti d'imposta e applicano una ritenuta del 21% (ventuno per cento) "
  "sull'ammontare dei canoni e corrispettivi all'atto del pagamento ai Mandanti, nei confronti dei locatori "
  "non esercenti attività d'impresa; conseguentemente, l'importo netto accreditato sul conto designato è già "
  "al netto di detta ritenuta, che costituisce — a decorrere dal 1° gennaio 2024 e indipendentemente dal "
  "regime fiscale (ordinario o cedolare secca) prescelto da ciascun Mandante, ai sensi dell'art. 4 c. 5 del "
  "D.L. 50/2017 come modificato dall'art. 1 c. 63 lett. b) della L. 30/12/2023 n. 213 — ritenuta a titolo di "
  "acconto delle imposte sui redditi, che ciascun Mandante è tenuto a scomputare nella propria dichiarazione "
  "dei redditi versando l'eventuale saldo entro i termini ordinari. La ritenuta è documentata dalla "
  "Certificazione Unica (CU) annuale rilasciata dal medesimo OTA. La Quota di cui all'art. 7.1 resta in ogni "
  "caso calcolata sul corrispettivo lordo del pernottamento e non sull'importo netto ai Mandanti. Resta ferma "
  "la diversa disciplina applicabile ove uno dei Mandanti sia esercente attività d'impresa, nel qual caso la "
  "ritenuta del 21% non si applica ed il Mandante interessato è tenuto a tutti gli ordinari adempimenti "
  "fiscali sui canoni percepiti, manlevando il Mandatario da ogni responsabilità conseguente a errata o "
  "omessa qualificazione del proprio regime fiscale;")

p(doc,
  "b) entro la fine di ciascun mese, i Mandanti si impegnano a corrispondere al Mandatario la quota maturata "
  "nel mese di riferimento, dietro presentazione di regolare fattura;")

p(doc,
  "c) in caso di ritardato pagamento di una fattura oltre 15 (quindici) giorni dall'emissione, considerato "
  "che entrambi i Mandanti agiscono in qualità di consumatori ai sensi dell'art. 5.3, sul debito scaduto "
  "decorrono di diritto gli interessi al saggio legale ex art. 1284 c.c., senza ulteriori maggiorazioni; il "
  "ritardo nei pagamenti reiterato per tre o più mensilità costituisce grave inadempimento e legittima il "
  "Mandatario alla risoluzione del Contratto ai sensi dell'art. 1454 c.c., previa diffida ad adempiere con "
  "termine non inferiore a 15 (quindici) giorni. Non opera la sospensione automatica dei Servizi né "
  "l'aggravamento del tasso al 1% mensile;")

p(doc,
  "d) in caso di risoluzione ex precedente lettera c), il Mandatario ha facoltà alternativa di: (i) "
  "continuare l'esecuzione dei Servizi sui soggiorni degli Ospiti già confermati, subordinatamente al "
  "pagamento anticipato della Quota stimata sui medesimi; ovvero (ii) cessare immediatamente l'esecuzione, "
  "fermi il diritto al risarcimento del danno e l'obbligo dei Mandanti di pagare la Quota dovuta sui "
  "soggiorni già confermati e svolti fino alla data di cessazione.")

p(doc,
  "7.4. Imposta di soggiorno. Le Parti concordano che, in tutti i Comuni in cui il regolamento comunale "
  "consenta la riscossione diretta da parte della piattaforma OTA, la riscossione dell'imposta di soggiorno "
  "avviene esclusivamente per il tramite dell'OTA, in modo che la piattaforma assuma la qualifica di soggetto "
  "responsabile del versamento ai sensi dell'art. 4 c. 5-ter D.L. 50/2017 (introdotto dall'art. 180 D.L. "
  "19/5/2020 n. 34 conv. L. 17/7/2020 n. 77 e oggetto di interpretazione autentica con efficacia retroattiva "
  "da parte dell'art. 5-quinquies D.L. 21/10/2021 n. 146 conv. L. 17/12/2021 n. 215), fermi restando gli "
  "eventuali obblighi dichiarativi, di comunicazione o di rendicontazione che il regolamento comunale ponga "
  "comunque in capo al gestore o al titolare della struttura ricettiva indipendentemente dalla modalità di "
  "riscossione, ai quali il Mandatario adempie per conto e nell'interesse dei Mandanti con il supporto "
  "informativo dei Mandanti medesimi. Solo nei Comuni in cui tale modalità non sia consentita dal "
  "regolamento, il Mandatario procede alla riscossione tramite POS al check-in o, in via residuale, in "
  "contanti, riversandola ai Mandanti con cadenza trimestrale unitamente a rendiconto dettagliato. Le Parti "
  "danno espressamente atto che, in questi ultimi Comuni, il Mandatario assume ex lege la qualifica di "
  "\"gestore\" tenuto agli obblighi dichiarativi e di versamento nei confronti del Comune, e che la manleva "
  "di cui infra opera nei soli rapporti interni tra le Parti e non è opponibile al Comune competente. I "
  "Mandanti sono e restano in ogni caso i soggetti responsabili degli obblighi dichiarativi e di versamento "
  "dell'imposta di soggiorno nei confronti del Comune competente, nei termini di legge, e si obbligano in "
  "solido a manlevare e tenere indenne il Mandatario da ogni sanzione, interesse o pretesa derivante da "
  "omessa, tardiva o incompleta dichiarazione o versamento dell'imposta, salvo il caso di dolo o colpa grave "
  "del Mandatario nella fase di raccolta e rendicontazione. Le Parti danno atto che gli specifici obblighi e "
  "responsabilità in materia possono variare in funzione del regolamento comunale applicabile, che i "
  "Mandanti sono tenuti a verificare.")

p(doc,
  "7.5. Il Mandatario redige e invia ai Mandanti, su base mensile, un resoconto dettagliato comprensivo di: "
  "prenotazioni ricevute, Quote di Pernottamento incassate dai Mandanti, costi sostenuti per conto dei "
  "Mandanti (servizi di pulizia, biancheria, manutenzioni, ricambi), quota dovuta al Mandatario.")

p(doc,
  "7.6. Il Mandatario, nello svolgimento dei Servizi relativi alla gestione dei prezzi, delle disponibilità "
  "e delle prenotazioni, si avvale di software di gestione (PMS / channel manager) di terze parti. Il costo "
  "mensile del software è pari a Euro 10,00 (dieci/00) IVA compresa per singolo Immobile. Per i primi 12 "
  "(dodici) mesi decorrenti dalla data di Pubblicazione dell'Annuncio, il costo del software è interamente a "
  "carico del Mandatario e non viene addebitato ai Mandanti. A partire dal 13° (tredicesimo) mese di vigenza "
  "del Contratto, il costo è a carico dei Mandanti e viene compensato direttamente sul saldo della "
  "rendicontazione mensile a favore dei Mandanti; ove la rendicontazione del mese di riferimento risulti "
  "incapiente o negativa, il costo è oggetto di fattura del Mandatario, regolata dai Mandanti mediante "
  "bonifico bancario entro 30 (trenta) giorni dalla data della fattura medesima. Il Mandatario non è "
  "autorizzato a conservare, trattare o movimentare dati di pagamento personali dei Mandanti (PAN di carta "
  "di credito, dati di addebito diretto SDD, credenziali di home banking) al di fuori dei canali ufficiali "
  "dei Fornitori di Servizi di Prenotazione Approvati e degli istituti di pagamento autorizzati. Qualora i "
  "Mandanti, alla scadenza dei primi 12 mesi, non accettino l'addebito del costo del software, il Mandatario "
  "può alternativamente (i) cessare l'uso del channel manager limitando la pubblicazione dell'Annuncio a un "
  "solo Fornitore di Servizi di Prenotazione Approvato; ovvero (ii) risolvere il Contratto ai sensi dell'art. "
  "1456 c.c. previa comunicazione scritta ai Mandanti con preavviso di 30 (trenta) giorni.")

p(doc,
  "7.7. Modifica della Quota. Il Mandatario può proporre per iscritto ai Mandanti una modifica della Quota di "
  "cui all'art. 7.1 esclusivamente per giustificato motivo, intendendosi per tale: (i) variazioni del regime "
  "IVA, di imposte o di oneri fiscali applicabili al Mandatario; (ii) adeguamento all'indice ISTAT FOI dei "
  "prezzi al consumo per le famiglie di operai e impiegati; (iii) documentato incremento dei costi dei "
  "servizi resi da terzi (channel-manager, pulizie, OTA fees, manutentori) superiore al 10% rispetto al "
  "momento della sottoscrizione; (iv) modifiche normative o regolamentari sopravvenute che incidano in modo "
  "significativo e documentabile sui costi di gestione. La proposta, motivata e accompagnata dalla "
  "documentazione del giustificato motivo, è comunicata ai Mandanti con preavviso di 90 (novanta) giorni "
  "rispetto alla data di efficacia. In assenza di accettazione scritta di entrambi i Mandanti entro 30 "
  "(trenta) giorni dal ricevimento della proposta, la modifica non ha effetto. In tale ipotesi, resta facoltà "
  "del Mandatario di recedere dal Contratto ai sensi dell'art. 8.2 con preavviso di 60 (sessanta) giorni; "
  "analoga facoltà di recesso spetta ai Mandanti, senza applicazione della penale di cui all'art. 8.2-bis "
  "anche qualora intervenga entro 6 mesi dalla Pubblicazione dell'Annuncio, con preavviso ridotto a 30 "
  "(trenta) giorni dalla data di ricevimento della proposta non accettata.")

p(doc,
  "7.8. Le modalità di pagamento di eventuali Servizi aggiuntivi concordati ai sensi dell'art. 2.2 saranno "
  "disciplinate caso per caso nel relativo accordo scritto.")

# --- ART 8 ---
h1(doc, "8. DURATA E CESSAZIONE")

p(doc,
  "8.1. Il presente Contratto ha durata di 12 (dodici) mesi, decorrenti dalla data di Pubblicazione "
  "dell'Annuncio, e si rinnova automaticamente alla scadenza per successivi periodi di 12 (dodici) mesi, "
  "salvo disdetta da farsi pervenire da una Parte all'altra mediante lettera raccomandata con avviso di "
  "ricevimento o posta elettronica certificata (PEC) all'indirizzo indicato all'art. 9, con preavviso di 60 "
  "(sessanta) giorni rispetto a ciascuna scadenza. La disdetta da parte dei Mandanti deve essere sottoscritta "
  "e trasmessa da entrambi congiuntamente. In pendenza del preavviso, i Servizi continuano a essere prestati "
  "regolarmente.")

p(doc,
  "8.2. Ciascuna Parte può recedere anticipatamente dal Contratto in qualunque momento mediante dichiarazione "
  "di recesso da inviare all'altra Parte con raccomandata A/R o PEC, almeno 60 (sessanta) giorni prima della "
  "data di efficacia. Il recesso dei Mandanti deve essere sottoscritto congiuntamente da entrambi.")

p(doc,
  "8.2-bis. Penale per recesso anticipato dei Mandanti entro 6 mesi. In deroga all'art. 8.2 e ai sensi "
  "dell'art. 1373 c.c., qualora i Mandanti esercitino il recesso anticipato prima del decorso di 6 (sei) mesi "
  "dalla data di Pubblicazione dell'Annuncio, gli stessi sono tenuti a corrispondere al Mandatario in solido, "
  "a titolo di rimborso forfettario delle spese di setup tecnico e operativo per la creazione dell'Annuncio "
  "(studio di mercato, redazione testi, sessione fotografica iniziale, configurazione Account OTA e channel "
  "manager, predisposizione operativa per accoglienza ospiti, briefing iniziale con imprese di pulizia e "
  "biancheria), una penale di Euro 200,00 (duecento/00) una tantum, oltre IVA dove applicabile. La penale è "
  "fatturata dal Mandatario contestualmente al recesso.")

p(doc, "La penale di cui al presente articolo non è dovuta nei seguenti casi:")
bullet(doc, "recesso dei Mandanti in conseguenza di variazione della Quota ex art. 7.7;")
bullet(doc, "recesso dei Mandanti in conseguenza di cessione del Contratto ai sensi dell'art. 11.2 a soggetto non accettato;")
bullet(doc, "recesso dei Mandanti per giusta causa, intendendosi tale qualunque inadempimento grave del Mandatario alle obbligazioni di cui all'art. 3.4;")
bullet(doc, "recesso conseguente a sopravvenuta impossibilità della prestazione non imputabile ai Mandanti;")
bullet(doc, "recesso dei Mandanti per mancato raggiungimento della soglia di occupazione di cui all'art. 8.5.")

p(doc,
  "Resta fermo che, anche in caso di recesso entro 6 mesi, l'Annuncio, le foto, le descrizioni e ogni "
  "contenuto pubblicato sull'Account restano di titolarità congiunta dei Mandanti ai sensi dell'art. 4.5, "
  "senza necessità di ulteriori atti di trasferimento.")

p(doc, "8.3. In caso di cessazione del Contratto:")
p(doc,
  "a) il Mandatario è tenuto a (i) svolgere i Servizi per tutti i soggiorni degli Ospiti già confermati alla "
  "data di cessazione, (ii) riconsegnare l'Immobile nello stato di fatto in cui si trova al momento della "
  "cessazione, (iii) trasferire ai Mandanti le credenziali e i materiali eventualmente in suo possesso, e "
  "cessare ogni operatività in qualità di Co-Host;")
p(doc,
  "b) i Mandanti sono tenuti in solido a (i) pagare al Mandatario quanto dovuto per i Servizi resi fino "
  "all'effettiva cessazione, e (ii) revocare tempestivamente i permessi di Co-Host del Mandatario sulle "
  "piattaforme.")

p(doc,
  "8.4. Il Mandatario potrà risolvere il presente Contratto ai sensi dell'art. 1456 c.c. in caso di "
  "inadempimento dei Mandanti a una o più delle obbligazioni di cui agli artt. 3.2, 3.3, 3.5, 4.1 e 7.3. "
  "Specularmente, i Mandanti potranno risolvere il Contratto ai sensi dell'art. 1456 c.c. in caso di "
  "reiterato e grave inadempimento del Mandatario alle obbligazioni di cui all'art. 3.4.")

p(doc,
  "8.5. Clausola di performance. Le Parti convengono che, qualora il tasso di occupazione medio dell'Immobile "
  "nei primi 6 (sei) mesi dalla data di Pubblicazione dell'Annuncio sia inferiore al 40% (quaranta per cento) "
  "— calcolato come rapporto tra notti effettivamente prenotate e notti disponibili nel medesimo arco "
  "temporale, escludendo dal denominatore i periodi bloccati dai Mandanti ai sensi dell'art. 3.2(e) e i "
  "giorni di indisponibilità tecnica imputabile ai Mandanti — i Mandanti potranno recedere dal presente "
  "Contratto con un preavviso di 30 (trenta) giorni, senza pagamento di alcuna penale e senza applicazione "
  "dell'art. 8.2-bis. Il calcolo del tasso di occupazione è effettuato sulla base del registro prenotazioni "
  "del channel-manager e del registro di diligenza di cui all'art. 6.5, ed è comunicato in trasparenza ai "
  "Mandanti alla scadenza del sesto mese.")

# --- ART 9 ---
h1(doc, "9. COMUNICAZIONI")

p(doc,
  "9.1. Qualsiasi comunicazione, istruzione o autorizzazione prevista dal presente Contratto deve essere "
  "effettuata, salvo diversa indicazione, ai recapiti di seguito indicati (o ai diversi recapiti che le Parti "
  "comunicheranno reciprocamente) tramite (i) posta elettronica certificata (PEC) o (ii) lettera raccomandata "
  "con avviso di ricevimento. Ai fini del presente Contratto, i Mandanti designano il Mandante 1 (Andrea "
  "Furlan) quale referente operativo principale per le comunicazioni ordinarie con il Mandatario, ferma "
  "restando la necessità di notifica ad entrambi i Mandanti per gli atti che comportino modifica, recesso o "
  "risoluzione del Contratto.")

p(doc, f'Recapiti del Mandante 1 (Andrea Furlan): indirizzo {MANDANTE1["residenza"]}.')
p(doc, f'Recapiti del Mandante 2 (Ephanie Sacramento): indirizzo {MANDANTE2["residenza"]}.')
p(doc, f'Recapiti del Mandatario: PEC {MANDATARIO["pec"]}; domicilio {MANDATARIO["domicilio"]}.')

# --- ART 10 ---
h1(doc, "10. PREVALENZA, MODIFICHE")

p(doc,
  "10.1. Il presente Contratto sostituisce e supera tutti i precedenti impegni, accordi, promesse, proposte, "
  "dichiarazioni, lettere di intenti, corrispondenza, comunicazioni — verbali o scritti — tra le Parti "
  "relativi all'oggetto del Contratto. Non esistono patti orali non completamente descritti nel presente "
  "Contratto.")

p(doc,
  "10.2. Nessuna modifica, rinuncia o esonero di responsabilità sarà efficace o vincolante per le Parti se "
  "non espressa per iscritto e con riferimento specifico al presente Contratto, e con la sottoscrizione di "
  "entrambi i Mandanti (oltre che del Mandatario).")

# --- ART 11 ---
h1(doc, "11. CESSIONE DEL CONTRATTO")

p(doc,
  "11.1. Il presente Contratto e i singoli diritti e obbligazioni da esso derivanti non possono essere ceduti "
  "o trasferiti a terzi, nemmeno per effetto di legge, senza il preventivo consenso scritto dell'altra Parte "
  "(o, per parte dei Mandanti, di entrambi congiuntamente), a pena di nullità.")

p(doc,
  "11.2. In deroga a quanto previsto all'art. 11.1, il Mandatario potrà cedere il presente Contratto previa "
  "comunicazione scritta ai Mandanti con preavviso di 60 (sessanta) giorni, nei seguenti casi: (i) "
  "costituzione di una società (a titolo esemplificativo S.r.l. o S.r.l.s.) per l'esercizio dell'attività di "
  "property management sotto il marchio Host Como, con cessione del Contratto a tale società; (ii) operazioni "
  "straordinarie (fusione, scissione, cessione d'azienda o di ramo d'azienda). Resta salva la facoltà dei "
  "Mandanti di recedere dal Contratto, senza preavviso e senza applicazione della penale di cui all'art. "
  "8.2-bis, qualora non accettino la cessione, entro 60 (sessanta) giorni dal ricevimento della "
  "comunicazione.")

# --- ART 12 ---
h1(doc, "12. INVALIDITÀ PARZIALE")

p(doc,
  "12.1. L'invalidità o inefficacia di qualsiasi disposizione del presente Contratto non comporta "
  "l'invalidità dell'intero Contratto. Le Parti si impegnano a negoziare in buona fede al fine di sostituire "
  "le disposizioni invalide con altre che producano, per quanto legalmente possibile, sostanzialmente gli "
  "stessi effetti.")

# --- ART 13 ---
h1(doc, "13. RINUNCIA")

p(doc,
  "13.1. Il mancato o ritardato esercizio da parte di una Parte di un diritto, potere o facoltà non "
  "costituisce rinuncia agli stessi. L'eventuale concessione di proroghe o dilazioni non modifica le "
  "responsabilità di ciascuna Parte.")

# --- ART 14 ---
h1(doc, "14. TUTELA DELLA SALUTE E SICUREZZA SUL LAVORO")

p(doc,
  "14.1. Le Parti danno atto che è stata omessa la redazione del DUVRI di cui all'art. 26, c. 3, D.Lgs. "
  "81/2008, in quanto i Servizi sono erogati in assenza di interferenze lavorative e non vi sono costi della "
  "sicurezza derivanti da interferenze.")

# --- ART 15 ---
h1(doc, "15. PROTEZIONE DEI DATI PERSONALI (GDPR)")

p(doc,
  "15.1. Le Parti si conformano al Regolamento (UE) 2016/679 (GDPR), al D.Lgs. 196/2003 (come modificato dal "
  "D.Lgs. 101/2018) e a ogni ulteriore normativa applicabile in materia di protezione dei dati personali.")

p(doc,
  "15.2. Ai fini del trattamento dei dati personali degli Ospiti necessari all'esecuzione del presente "
  "Contratto, i Mandanti sono Titolari del trattamento; il Mandatario è Responsabile del trattamento ai sensi "
  "dell'art. 28 GDPR, in quanto tratta i dati per conto dei Mandanti nell'ambito dei Servizi, avvalendosi di "
  "sub-responsabili ai sensi dell'art. 28.2 GDPR secondo quanto previsto nell'Allegato D.")

p(doc,
  "15.3. Il trattamento è disciplinato in dettaglio dal Data Processing Agreement (DPA) di cui all'Allegato "
  "D, che costituisce parte integrante del presente Contratto.")

p(doc,
  "15.4. I dati personali delle Parti contraenti sono trattati per finalità di esecuzione del Contratto, "
  "adempimento di obblighi di legge (fiscali e di fatturazione), gestione del rapporto contrattuale e — "
  "previo specifico consenso — finalità di marketing, ai sensi degli artt. 6.1(b), 6.1(c) e 6.1(a) GDPR. I "
  "dati sono conservati per il tempo necessario alle predette finalità e comunque per i termini prescritti "
  "dalla normativa fiscale (10 anni ai sensi dell'art. 2220 c.c.).")

p(doc,
  "15.5. Le Parti possono esercitare in qualsiasi momento i diritti di cui agli artt. 15-22 GDPR (accesso, "
  "rettifica, cancellazione, limitazione, portabilità, opposizione) scrivendo agli indirizzi indicati "
  "all'art. 9.")

# --- ART 16 ---
h1(doc, "16. NUMERO DI ORIGINALI, REGISTRAZIONE, IMPOSTA DI BOLLO")

p(doc,
  "16.1. Il presente Contratto viene sottoscritto in n. 3 (tre) esemplari originali, uno per ciascuna delle "
  "Parti (Mandante 1, Mandante 2, Mandatario). Il Contratto non è soggetto all'obbligo di registrazione, "
  "salvo il caso d'uso ai sensi del DPR 131/1986. L'eventuale imposta di bollo è a carico dei Mandanti in "
  "solido.")

# --- ART 17 ---
h1(doc, "17. LEGGE APPLICABILE E FORO COMPETENTE")

p(doc,
  "17.1. Il presente Contratto è disciplinato dalla legge italiana, con esclusione delle norme di diritto "
  "internazionale privato.")

p(doc,
  "17.2. Per tutte le controversie concernenti l'interpretazione, validità, efficacia, esecuzione o "
  "scioglimento del presente Contratto sarà competente in via esclusiva il Foro di Como.")

p(doc,
  "17.3. Resta espressamente salvo il foro inderogabile del consumatore di cui all'art. 33, comma 2, lett. u) "
  "del Codice del Consumo (D.Lgs. 206/2005), qualora uno dei Mandanti sia qualificabile come \"consumatore\" "
  "ai sensi dell'art. 3 del medesimo Codice: in tal caso, sarà competente il foro del luogo di residenza o "
  "domicilio elettivo del Mandante consumatore che agisce o è convenuto in giudizio.")

# --- SOTTOSCRIZIONI ---
p(doc, "", space_after=12)
p(doc, "Luogo e data: ______________________________, lì _____ / _____ / __________")
p(doc, "", space_after=18)

def firma_block(doc, label):
    p(doc, "_____________________________________________________")
    p(doc, label, bold=True)
    p(doc, "", space_after=12)

firma_block(doc, "Il Mandante 1 (Andrea Furlan)")
firma_block(doc, "Il Mandante 2 (Ephanie Sacramento)")
firma_block(doc, "Il Mandatario (Angelo Talarico)")

# --- ACCETTAZIONE CLAUSOLE VESSATORIE ---
h1(doc, "Accettazione specifica di clausole vessatorie (artt. 1341 e 1342 c.c.)")

p(doc,
  "Ai sensi e per gli effetti degli artt. 1341 e 1342 c.c. — ferma la disciplina inderogabile di cui all'art. "
  "5.3 ove applicabile — i Mandanti dichiarano di approvare specificamente, previa attenta lettura, le "
  "seguenti clausole: art. 2.2 (esclusione dell'art. 1661 c.c. e facoltà di rifiuto di servizi aggiuntivi); "
  "art. 2.4 (esclusione dei soggiorni superiori a 30 giorni dall'oggetto del mandato); art. 2.5 (facoltà di "
  "subappalto e avvalimento di terzi); art. 3.2 (obblighi a carico dei Mandanti); art. 3.3 (clausola "
  "risolutiva espressa in caso di cancellazione prenotazioni); art. 3.5 (esclusiva); art. 3.6 (conseguenze "
  "del mancato ottenimento del CIN, sospensione automatica dell'Annuncio e risoluzione ex art. 1456 c.c. "
  "decorsi 90 giorni); art. 3.8 (checklist di sicurezza all'onboarding ed efficacia probatoria); art. 4.2 "
  "(discrezionalità nei canali OTA); art. 5.1, lett. a) (esonero dall'obbligo di consultazione preventiva); "
  "art. 5.2 (manleva dei Mandanti, limitata ai fatti di propria sfera di controllo); art. 6 (limitazioni di "
  "responsabilità; sistema integrato di copertura ex art. 6.4; registro di diligenza ex art. 6.5 con "
  "efficacia probatoria); art. 7.1 (Quota base, incremento di 2,5 punti sotto soglia 150 giorni, calcolo sul "
  "lordo); art. 7.1-bis (applicazione della Quota alle sole prenotazioni dirette per le quali il Mandatario "
  "rende effettivamente i Servizi; esenzione delle prenotazioni gestite in autonomia dai Mandanti; eccezioni "
  "per soggiorni gratuiti personali dei Mandanti e dei familiari nei limiti di 30 giorni/anno); art. 7.1-ter "
  "(autorizzazione ai Mandanti a creare un proprio sito web dell'Immobile con incassi via PSP; Quota del 10% "
  "sempre dovuta sul corrispettivo netto); art. 7.3 (flusso pagamenti e interessi di mora); art. 7.4 "
  "(responsabilità dei Mandanti per dichiarazione/versamento imposta di soggiorno e relativa manleva); art. "
  "7.6 (channel manager a carico dei Mandanti dal 13° mese e facoltà alternative del Mandatario in caso di "
  "rifiuto); art. 7.7 (modifica della Quota soltanto per giustificato motivo e diritto di recesso senza "
  "penale); art. 8.1 (rinnovo automatico); art. 8.2-bis (penale di Euro 200,00 per recesso anticipato dei "
  "Mandanti entro 6 mesi, con esclusioni); art. 8.4 (clausole risolutive espresse); art. 10.2 (forma scritta "
  "ad substantiam); art. 11 (limiti alla cessione del Contratto); art. 13.1 (esclusione di rinuncia per "
  "inerzia); art. 17.2 (Foro di Como, fatto salvo l'art. 17.3 in materia di foro inderogabile del "
  "consumatore).")

p(doc, "", space_after=18)
firma_block(doc, "Il Mandante 1 (Andrea Furlan) — per accettazione specifica")
firma_block(doc, "Il Mandante 2 (Ephanie Sacramento) — per accettazione specifica")

# --- ALLEGATO A ---
doc.add_page_break()
h1(doc, "ALLEGATO A — DESCRIZIONE DEI SERVIZI")

p(doc,
  "I Servizi di seguito descritti sono resi dal Mandatario, il quale può eseguirli direttamente ovvero, sotto "
  "la propria esclusiva responsabilità e secondo quanto previsto all'art. 2.5, avvalendosi di operatori, "
  "fornitori e prestatori di servizi autonomi. In ogni caso il Mandatario resta l'unico soggetto obbligato "
  "verso i Mandanti.")

p(doc, "Sezione I — Servizi operativi e di accoglienza", bold=True)

p(doc,
  "A.I.1. Coordinamento dei Servizi Domestici. Prima dell'arrivo di ciascun Ospite, il Mandatario coordina "
  "l'esecuzione dei seguenti servizi, avvalendosi di imprese terze specializzate (i cui costi sono fatturati "
  "direttamente ai Mandanti): (i) rimozione di rifiuti, stoviglie e biancheria sporca lasciati dagli Ospiti "
  "precedenti; (ii) pulizia con detergenti idonei di tutte le superfici; (iii) rifacimento dei letti con "
  "biancheria pulita; (iv) aspirazione tappeti e lavaggio pavimenti; (v) verifica del funzionamento di "
  "apparecchiature elettriche, elettrodomestici e dispositivi di sicurezza, con segnalazione tempestiva di "
  "malfunzionamenti; (vi) riassortimento di beni di prima necessità (carta igienica, sapone mani, "
  "bagnoschiuma, shampoo, sale lavastoviglie, sacchi spazzatura); (vii) predisposizione della biancheria "
  "pulita per ciascuna prenotazione (n. 1 lenzuolo pulito/stirato per letto, copri piumone, federe in numero "
  "adeguato, n. 1 asciugamano bagno + n. 1 asciugamano viso per Ospite, n. 1 tappetino bagno per bagno).")

p(doc,
  "A.I.2. Accoglienza Ospiti. Il Mandatario o suo incaricato provvede a (i) accogliere personalmente ciascun "
  "Ospite all'arrivo, ovvero, ove concordato, predisporre check-in self-service tramite smart lock con "
  "istruzioni dettagliate; (ii) consegnare le chiavi (o codici di accesso) e illustrare il funzionamento "
  "dell'Immobile; (iii) fornire informazioni utili sulla zona (POI, ristoranti, trasporti, contatti "
  "d'emergenza).")

p(doc,
  "A.I.3. Comunicazione operativa con gli Ospiti durante il soggiorno. Il Mandatario gestisce in autonomia "
  "ogni richiesta degli Ospiti durante il soggiorno: assistenza, emergenze, sostituzione asciugamani, "
  "malfunzionamenti, smarrimento chiavi, problemi tecnici risolvibili sul posto.")

p(doc,
  "A.I.4. Coordinamento manutenzioni e segnalazione danni. Il Mandatario (i) coordina gli interventi di "
  "manutenzione ordinaria e straordinaria, gestendo direttamente o tramite tecnici specializzati (idraulici, "
  "elettricisti, antennisti); i costi di ricambi, materiali e prestazioni di tecnici sono a carico dei "
  "Mandanti e fatturati separatamente, con preventiva approvazione dei Mandanti per interventi superiori a "
  "Euro 100,00 (cento/00); (ii) in caso di intervento urgente di importo inferiore a Euro 50,00 durante un "
  "soggiorno Ospite, procede autonomamente per evitare disservizi, addebitando il costo nel resoconto "
  "mensile; (iii) in caso di intervento di importo superiore durante un soggiorno Ospite, richiede il "
  "preventivo consenso scritto dei Mandanti; (iv) riscontra e segnala tempestivamente eventuali danni "
  "all'Immobile al termine di ciascun soggiorno, attivando ove possibile i rimborsi previsti dalla "
  "piattaforma (es. AirCover di Airbnb) o ricorrendo al deposito cauzionale eventualmente trattenuto "
  "sull'Ospite.")

p(doc,
  "A.I.5. Adempimenti normativi operativi per conto dei Mandanti. Il Mandatario, per conto dei Mandanti e "
  "utilizzando le credenziali di questi ultimi, provvede a: (i) Alloggiati Web — registrazione di ciascun "
  "Ospite sul portale della Polizia di Stato ai sensi dell'art. 109 TULPS entro 24 (ventiquattro) ore "
  "dall'arrivo dell'Ospite, ovvero entro le 6 (sei) ore successive all'arrivo qualora il soggiorno abbia "
  "durata inferiore alle 24 ore; (ii) Comunicazioni ISTAT — invio mensile dei dati di flusso turistico "
  "secondo le modalità previste dalla Regione Lombardia e dal Comune competente; (iii) Imposta di soggiorno "
  "— riscossione dagli Ospiti secondo le modalità di cui all'art. 7.4, tracciamento separato delle somme "
  "raccolte, consegna trimestrale ai Mandanti con rendiconto dettagliato per il riversamento al Comune nei "
  "termini di legge.")

p(doc,
  "Obblighi dei Mandanti connessi ai Servizi operativi. I Mandanti si impegnano a: (a) tenere a disposizione "
  "presso l'Immobile aspirapolvere, attrezzature per la pulizia e materiali di consumo necessari; (b) "
  "informare di eventuali oggetti particolarmente fragili o di valore, contrassegnandoli; (c) custodire in "
  "armadio chiuso a chiave eventuali oggetti di valore o riservati; (d) fornire al Mandatario le credenziali "
  "del portale Alloggiati Web e i codici di accesso all'imposta di soggiorno comunale.")

p(doc, "Sezione II — Servizi di gestione commerciale, tecnologica e amministrativa", bold=True)

p(doc,
  "A.II.1. Studio di mercato, Pricing strategy e Dynamic pricing. Il Mandatario (i) effettua uno studio dei "
  "prezzi di mercato applicabili all'Immobile, basato sull'analisi della stagionalità, degli eventi in corso "
  "sul territorio del Lago di Como, della zona, delle caratteristiche dell'Immobile e dei competitor; (ii) "
  "implementa dynamic pricing con variazione continua delle tariffe in funzione di domanda, eventi locali, "
  "stagionalità, lead-time prenotazione, performance storica dell'Immobile; (iii) monitora la performance e "
  "ottimizza le tariffe in modo iterativo.")

p(doc,
  "A.II.2. Creazione e gestione tecnica dell'Annuncio. Il Mandatario (i) crea per conto dei Mandanti "
  "l'annuncio sull'Account dei Fornitori di Servizi di Prenotazione Approvati (in via iniziale: Airbnb, "
  "Booking.com, Expedia; ulteriori canali a sua discrezione secondo l'art. 4.2); (ii) ottimizza l'annuncio "
  "(titolo, descrizione, foto, tag, posizionamento) in funzione delle migliori pratiche di conversione di "
  "ciascun canale; (iii) mantiene il calendario sempre aggiornato e sincronizzato tra i canali tramite "
  "channel manager; (iv) gestisce le recensioni rispondendo secondo le best practice di ospitalità.")

p(doc,
  "A.II.3. Gestione del software PMS / Channel Manager. Il Mandatario provvede all'installazione, "
  "configurazione e gestione operativa del software PMS e channel manager necessario alla gestione "
  "multi-canale delle prenotazioni e dei prezzi. Il costo mensile del software, pari a Euro 10,00 IVA "
  "compresa per Immobile, è interamente a carico del Mandatario per i primi 12 (dodici) mesi dalla data di "
  "Pubblicazione dell'Annuncio e successivamente a carico dei Mandanti, secondo quanto previsto all'art. 7.6. "
  "Il Mandatario può configurare sui canali OTA un markup tariffario volto a neutralizzare le service fee a "
  "carico dell'host, in modo da garantire ai Mandanti un payout netto coerente con i prezzi base impostati.")

p(doc,
  "A.II.4. Reportistica e Rendicontazione. Il Mandatario redige e invia ai Mandanti, su base mensile, un "
  "resoconto dettagliato comprensivo di: prenotazioni ricevute (sia tramite OTA sia da canali diversi, "
  "inclusi i contatti privati e diretti dei Mandanti o di terzi, ai sensi dell'art. 7.1-bis); Quote di "
  "Pernottamento maturate (incassate dai Mandanti per il tramite OTA ovvero concordate dai Mandanti con "
  "l'Ospite per le prenotazioni dirette); costi sostenuti per conto dei Mandanti (servizi di pulizia, "
  "biancheria, manutenzioni, ricambi); quota dovuta al Mandatario, con evidenza separata delle componenti "
  "riferite a prenotazioni dirette e di eventuali soggiorni gratuiti di familiari ex art. 7.1-bis.")

p(doc,
  "A.II.5. Consulenza commerciale ai Mandanti. Il Mandatario fornisce ai Mandanti consulenza gestionale e "
  "commerciale in tema di scelta dei canali OTA da attivare, strategie di posizionamento online, "
  "ottimizzazione dell'occupancy, valutazioni di marketing turistico.")

p(doc,
  "A.II.6. Supporto agli adempimenti amministrativi. Il Mandatario fornisce ai Mandanti supporto operativo "
  "per (i) la richiesta del CIN presso la BDSR del MITUR ai sensi dell'art. 13-ter D.L. 145/2023 conv. L. "
  "191/2023, del D.M. MITUR 6/6/2024 e dell'Avviso MITUR pubblicato in GU n. 130 P. II del 3/9/2024 "
  "(obblighi e sanzioni applicabili dal 1° gennaio 2025 in forza della proroga disposta con avviso MITUR del "
  "22/10/2024); (ii) la richiesta del CIR Lombardia se ancora applicabile ai sensi della L.R. 27/2015 e del "
  "R.R. 5/8/2016 n. 7; (iii) la presentazione della SCIA o comunicazione di inizio attività al Comune di "
  "Cernobbio; (iv) l'orientamento informativo sulla scelta del regime fiscale applicabile ai Mandanti "
  "(cedolare secca 21%/26%; IRPEF ordinario; regime forfetario; regime IVA). Resta inteso che il Mandatario "
  "non sostituisce il commercialista dei Mandanti né presta consulenza fiscale qualificata, ma fornisce "
  "orientamento informativo e supporto operativo.")

# --- ALLEGATO B ---
doc.add_page_break()
h1(doc, "ALLEGATO B — PREREQUISITI DELL'IMMOBILE E DOTAZIONI OBBLIGATORIE DI SICUREZZA")

p(doc,
  "L'Immobile deve essere fornito di tutte le dotazioni di seguito elencate. È onere e cura dei Mandanti "
  "provvedere all'acquisto, installazione e manutenzione di tali dotazioni in conformità alla normativa "
  "vigente.")

p(doc,
  "B.1. Dotazioni di sicurezza obbligatorie. I Mandanti adottano e mantengono le dotazioni di sicurezza "
  "dell'Immobile in coerenza con la normativa in materia di sicurezza degli impianti (D.M. 22/01/2008 n. 37), "
  "con la disciplina regionale per le strutture ricettive non alberghiere (R.R. Lombardia 5/8/2016 n. 7), "
  "con la disciplina antincendio nei limiti in cui sia applicabile all'Immobile in ragione delle sue "
  "caratteristiche e dimensione, nonché con le raccomandazioni di sicurezza dei principali Fornitori di "
  "Servizi di Prenotazione Approvati. Resta a esclusivo carico dei Mandanti la verifica della normativa "
  "effettivamente applicabile all'Immobile e l'adempimento degli obblighi che ne discendono. A titolo "
  "indicativo e non esaustivo, le Parti danno atto delle seguenti dotazioni opportune o richieste dagli OTA:")

DOT = [
    "rilevatore di monossido di carbonio (CO) nei locali con impianti a combustione (caldaie, stufe, caminetti, cucine a gas) — obbligatorio;",
    "rilevatore di fumo in ciascun piano dell'Immobile — fortemente raccomandato (obbligatorio in alcuni Comuni);",
    "estintore a polvere o CO₂ (almeno classe 21A 89B C, capacità minima 2 kg) — fortemente raccomandato, obbligatorio per case vacanze classificate ai sensi della normativa regionale Lombardia;",
    "coperta antifiamma nel locale cucina — raccomandato;",
    "cassetta di pronto soccorso completa e in corso di validità — raccomandato;",
    "numeri di emergenza chiaramente esposti (112, 118, 115, contatto del Mandatario, ospedale più vicino) — obbligatorio;",
    "planimetria di emergenza con indicazione delle vie di fuga — obbligatorio per strutture con più camere;",
    "dichiarazione di conformità degli impianti (elettrico, idraulico, gas) ai sensi del D.M. 37/2008 — i Mandanti devono detenerla e fornirla in copia su richiesta;",
    "certificazione di conformità della caldaia e libretto di impianto aggiornato;",
    "salvavita / differenziale dell'impianto elettrico funzionante e periodicamente testato — obbligatorio;",
    "serrature a doppia mandata alla porta d'ingresso principale — raccomandato.",
]
for x in DOT:
    bullet(doc, x)

p(doc,
  "I Mandanti dichiarano, sotto la propria responsabilità, di provvedere a installare e mantenere efficienti "
  "le suddette dotazioni e a custodire tutta la documentazione tecnica relativa, esibendola al Mandatario su "
  "richiesta.")

p(doc,
  "B.2. Stanza da letto: n. 1 copriletto/piumone/coperta per ciascun letto; n. 4 cuscini per ciascun letto "
  "matrimoniale (n. 2 per letto singolo); n. 1 coprimaterasso per ciascun letto; n. 1 appendiabiti per "
  "camera; n. 1 cestino per rifiuti.")

p(doc,
  "B.3. Soggiorno: divano comodo; libri / giochi da tavolo (raccomandati); TV con telecomando; riscaldamento "
  "e/o climatizzazione funzionanti.")

p(doc,
  "B.4. Bagno: n. 1 asciugacapelli; n. 1 porta spazzolino; n. 1 specchio (almeno uno per figura intera "
  "nell'Immobile); porta rotolo di carta igienica; scopino per water; cestino per la spazzatura.")

p(doc,
  "B.5. Cucina: tostapane; bollitore; posate, bicchieri, piatti, tazze e utensili in numero almeno doppio "
  "rispetto agli ospiti massimi; cestino per la spazzatura; forno a microonde; macchina del caffè (anche "
  "moka).")

p(doc,
  "B.6. Servizi e attrezzature comuni: WiFi a banda larga (almeno 30 Mbps download) con copertura in tutti "
  "gli ambienti; mocio con secchio; aspirapolvere; ferro e asse da stiro; stendino per biancheria; "
  "termosifoni e/o climatizzatori funzionanti e adeguatamente manutenuti.")

# --- ALLEGATO C ---
doc.add_page_break()
h1(doc, "ALLEGATO C — TARIFFE E CONDIZIONI ECONOMICHE")

p(doc,
  "C.1. Sono compresi nella commissione del 10% di cui all'art. 7.1: gestione operativa e di accoglienza "
  "(coordinamento dei servizi di pulizia, i cui costi restano a carico dei Mandanti; coordinamento dei "
  "servizi di biancheria; accoglienza Ospiti in fascia oraria standard 10:00-20:00; comunicazione operativa "
  "con gli Ospiti durante il soggiorno; coordinamento manutenzioni ordinarie e straordinarie; registrazione "
  "su Alloggiati Web; comunicazioni ISTAT mensili; raccolta dell'imposta di soggiorno con consegna "
  "trimestrale ai Mandanti); gestione commerciale, tecnologica e amministrativa (attivazione e gestione "
  "annuncio sui Fornitori di Servizi di Prenotazione Approvati — in via iniziale Airbnb, Booking.com, "
  "Expedia, ulteriori canali a discrezione del Mandatario; dynamic pricing e gestione calendario "
  "multi-canale; gestione tecnica del software PMS / channel manager; reportistica mensile e "
  "rendicontazione; consulenza commerciale ai Mandanti; supporto agli adempimenti amministrativi CIN, CIR, "
  "SCIA).")

p(doc,
  "C.2. Qualora l'Immobile sia reso disponibile per la locazione turistica per un periodo inferiore a 150 "
  "(centocinquanta) giorni nell'arco di ciascun anno contrattuale, la Quota sulle prenotazioni "
  "effettivamente realizzate è automaticamente incrementata di 2,5 (due virgola cinque) punti percentuali, "
  "passando dal 10% al 12,5% (dodici virgola cinque per cento) oltre IVA dove applicabile, a decorrere dal "
  "momento in cui la soglia viene mancata.")

p(doc,
  "C.3. I Mandanti che intendano ridurre i giorni di disponibilità devono comunicarlo al Mandatario con "
  "preavviso minimo di 15 (quindici) giorni ai sensi dell'art. 3.2(e), e in ogni caso non oltre periodi già "
  "oggetto di prenotazione confermata.")

p(doc,
  "C.4. I costi di riparazioni e manutenzione restano a carico dei Mandanti secondo quanto previsto "
  "all'art. 3.2(h) e all'Allegato A, sezione A.I.4.")

p(doc,
  "C.5. Le previsioni di rendimento eventualmente comunicate dal Mandatario ai Mandanti sono basate sullo "
  "studio del mercato turistico locale e sull'esperienza professionale del Mandatario. Tali stime sono "
  "meramente indicative e non costituiscono garanzia di risultato.")

p(doc,
  "C.6. Sono Servizi extra, da concordare di volta in volta per iscritto tra le Parti con specifica del "
  "corrispettivo, a titolo esemplificativo e non esaustivo: restyling o home staging dell'Immobile; shooting "
  "fotografico professionale; check-in fuori fascia oraria standard (prima delle 10:00 o dopo le 20:00, "
  "festivi); check-in con servizio multilingua aggiuntivo; traduzione manuale d'uso e guide turistiche "
  "personalizzate; deep cleaning straordinario; gestione lavanderia con frequenza superiore allo standard; "
  "installazione e gestione smart lock / serrature digitali; predisposizione welcome basket personalizzati.")

p(doc,
  "C.7. L'impresa di pulizie incaricata non fornisce, salvo specifico accordo aggiuntivo a tariffa "
  "concordata: lavaggio di muri e soffitti; trattamento professionale di mobili o tappeti; lavaggio esterno "
  "delle finestre; rimozione sporcizia causata da animali; giardinaggio e pulizia capanni; pulizia patii, "
  "terrazzi o giardini estesi; rimozione di muffe o trattamenti antimuffa; pulizie industriali; spostamento "
  "di mobili pesanti; pulizia di superfici inaccessibili da terra senza scala; pulizia di superfici "
  "fortemente incrostate; derattizzazioni o disinfestazioni; pulizia di cortili o garage.")

# --- ALLEGATO D ---
doc.add_page_break()
h1(doc, "ALLEGATO D — DATA PROCESSING AGREEMENT (DPA) EX ART. 28 GDPR")

p(doc,
  "D.1. Il presente DPA disciplina il trattamento dei dati personali degli Ospiti effettuato dal Mandatario "
  "(Responsabile del trattamento) per conto dei Mandanti (Titolari del trattamento) nell'ambito dei Servizi "
  "del Contratto.")

p(doc,
  "D.2. Oggetto e natura del trattamento. Categorie di dati trattati: dati anagrafici (nome, cognome, data "
  "di nascita, luogo di nascita, cittadinanza), dati di contatto (email, telefono), dati di documento di "
  "identità (tipo, numero, scadenza), metadati di transazione visualizzati dal Mandatario nel pannello "
  "Co-Host degli OTA (importo lordo, data, status payout; esclusi IBAN, numeri di carta o credenziali "
  "bancarie), dati di soggiorno (date, numero ospiti, preferenze comunicate). Categorie di interessati: "
  "Ospiti che effettuano o richiedono prenotazioni presso l'Immobile. Finalità del trattamento: esecuzione "
  "del Contratto, gestione della prenotazione e del soggiorno, adempimenti normativi (Alloggiati Web, "
  "ISTAT, imposta di soggiorno, fatturazione), comunicazioni operative con gli Ospiti, gestione recensioni. "
  "Durata del trattamento: per tutta la durata del Contratto e, successivamente, per i termini di legge.")

p(doc,
  "D.3. Obblighi del Mandatario quale Responsabile del trattamento. Il Mandatario si impegna a: (a) trattare "
  "i dati personali esclusivamente per le finalità del Contratto e nei limiti delle istruzioni documentate "
  "dei Mandanti; (b) garantire che i propri operatori, dipendenti, collaboratori e subappaltatori "
  "autorizzati al trattamento siano vincolati a un obbligo di riservatezza; (c) adottare le misure di "
  "sicurezza tecniche e organizzative adeguate ai sensi dell'art. 32 GDPR (cifratura dati in transito, "
  "controllo accessi, backup, registro trattamenti); (d) assistere i Mandanti nell'evasione delle richieste "
  "degli interessati (artt. 15-22 GDPR); (e) notificare ai Mandanti senza ingiustificato ritardo e comunque "
  "entro 48 (quarantotto) ore qualsiasi violazione di sicurezza che riguardi i dati personali oggetto del "
  "presente DPA; (f) al termine del Contratto, restituire ai Mandanti o cancellare tutti i dati personali "
  "trattati per suo conto, salvo conservazione imposta dalla legge; (g) tenere un registro delle attività di "
  "trattamento svolte per conto dei Mandanti (art. 30.2 GDPR); (h) non trasferire dati personali al di fuori "
  "dello Spazio Economico Europeo se non con garanzie adeguate.")

p(doc,
  "D.4. Sub-responsabili del trattamento. I Mandanti autorizzano il Mandatario, anche ai sensi dell'art. "
  "28.2 GDPR, ad avvalersi di sub-responsabili per l'esecuzione di specifiche attività di trattamento, "
  "espressamente compresi: fornitori di software PMS/channel manager, imprese di pulizia, lavanderia "
  "industriale, gestori OTA, nonché il prestatore di servizi autonomo Crapotca Andrei (P.IVA 04236480135) "
  "per le attività di gestione tecnica dell'Annuncio, gestione del PMS/channel manager, reportistica ai "
  "Mandanti e comunicazione pre-prenotazione con gli ospiti potenziali. Il Mandatario garantisce che i "
  "sub-responsabili siano vincolati da impegni di protezione dei dati equivalenti a quelli previsti dal "
  "presente DPA.")

p(doc,
  "Il Mandatario informa i Mandanti con preavviso scritto di almeno 30 (trenta) giorni di qualsiasi modifica "
  "alla lista dei sub-responsabili (aggiunta, sostituzione o revoca), indicando i dettagli del nuovo "
  "soggetto, l'area di trattamento affidata e una sintesi delle garanzie di sicurezza adottate. I Mandanti "
  "possono opporsi per giustificati motivi entro 15 (quindici) giorni dalla comunicazione; in tal caso le "
  "Parti negoziano in buona fede una soluzione alternativa, restando salva la facoltà dei Mandanti di "
  "recedere dal Contratto qualora non si raggiunga un accordo entro ulteriori 30 (trenta) giorni dalla data "
  "dell'opposizione.")

p(doc,
  "D.5. Cooperazione con l'Autorità di controllo. Il Mandatario coopera, su richiesta, con il Garante per la "
  "protezione dei dati personali e con le altre Autorità competenti.")

p(doc,
  "D.6. Responsabilità. Ciascuna Parte risponde dei danni derivanti dal trattamento dei dati personali solo "
  "per quanto attiene agli obblighi specificamente posti a proprio carico ai sensi del GDPR e del presente "
  "DPA.")

p(doc, "", space_after=18)
firma_block(doc, "Il Mandante 1 (Andrea Furlan) — per il DPA ex Allegato D")
firma_block(doc, "Il Mandante 2 (Ephanie Sacramento) — per il DPA ex Allegato D")
firma_block(doc, "Il Mandatario (Angelo Talarico) — per il DPA ex Allegato D")

# Salva
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"OK: {OUTPUT}")
print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")
